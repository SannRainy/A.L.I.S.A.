# -*- coding: utf-8 -*-
"""
conftest.py — Pytest Configuration & Automatic Test Report Generator (Word .docx & Markdown .md)
==============================================================================================
Otomatis membuat/menimpa file laporan pengujian di folder 'test_reports/':
  · test_reports/laporan.docx  (Microsoft Word — dengan tabel garis horizontal murni)
  · test_reports/laporan.md    (Markdown Format)
"""

import os
import sys
import time
from datetime import datetime
import pytest

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn


def _get_module_attr(mod_suffix: str, attr_name: str) -> list:
    """Cari atribut hasil di sys.modules dengan mencocokkan nama file/modul."""
    for k, mod in list(sys.modules.items()):
        if mod and (k == mod_suffix or k.endswith("." + mod_suffix)):
            if hasattr(mod, attr_name):
                return getattr(mod, attr_name)
    return []


def _set_cell_background(cell, fill_hex: str):
    """Set warna latar belakang sel (shading)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def _set_table_horizontal_borders_only(table, border_color="2B579A", inside_color="D3D3D3"):
    """
    Mengatur border tabel agar HANYA garis horizontal yang terlihat (Top, Bottom, InsideH).
    Garis vertikal (Left, Right, InsideV) disembunyikan total.
    """
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="12" w:space="0" w:color="{border_color}"/>
        <w:bottom w:val="single" w:sz="12" w:space="0" w:color="{border_color}"/>
        <w:left w:val="none"/>
        <w:right w:val="none"/>
        <w:insideH w:val="single" w:sz="6" w:space="0" w:color="{inside_color}"/>
        <w:insideV w:val="none"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))


def _set_table_cell_margins(table, top=120, bottom=120, left=150, right=150):
    """Set padding dalam sel tabel (dalam dxa, 20 dxa = 1 pt)."""
    tblPr = table._tbl.tblPr
    margins_xml = f'''
    <w:tblCellMar {nsdecls("w")}>
        <w:top w:w="{top}" w:type="dxa"/>
        <w:bottom w:w="{bottom}" w:type="dxa"/>
        <w:left w:w="{left}" w:type="dxa"/>
        <w:right w:w="{right}" w:type="dxa"/>
    </w:tblCellMar>
    '''
    tblPr.append(parse_xml(margins_xml))


def _format_word_cell(cell, text: str, bold=False, italic=False, color=None, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Helper untuk format teks di sel Word secara konsisten."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(str(text))
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color


def pytest_sessionfinish(session, exitstatus):
    """
    Hook pytest yang dipicu otomatis setelah seluruh pengujian selesai.
    Merender file 'test_reports/laporan.docx' (Word) & 'test_reports/laporan.md' (Markdown)
    dengan format per-skenario (5 kasus uji per tabel) dan tema warna gelap (Hitam 50%).
    """
    current_dir = os.path.dirname(os.path.abspath(__file__))        # backend/tests/layers
    tests_dir   = os.path.dirname(current_dir)                     # backend/tests
    backend_dir = os.path.dirname(tests_dir)                       # backend
    root_dir    = os.path.dirname(backend_dir)                      # project root

    report_dir  = os.path.join(root_dir, "test_reports")
    os.makedirs(report_dir, exist_ok=True)
    docx_file   = os.path.join(report_dir, "laporan.docx")
    md_file     = os.path.join(report_dir, "laporan.md")

    if root_dir not in sys.path:
        sys.path.insert(0, root_dir)

    try:
        import build_full_test_report
        now_str = datetime.now().strftime("%d %B %Y - %H:%M:%S WIB")
        build_full_test_report.generate_docx(docx_file, now_str)
        build_full_test_report.generate_md(md_file, now_str)
        return
    except Exception as err:
        print(f"\n[WARNING] Gagal memanggil generator build_full_test_report: {err}")


    # Page margins
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    COLOR_PRIMARY = RGBColor(27, 54, 93)     # Deep Navy #1B365D
    COLOR_SUCCESS = RGBColor(34, 139, 34)    # Forest Green
    COLOR_FAIL    = RGBColor(178, 34, 34)    # Firebrick Red
    COLOR_TEXT    = RGBColor(51, 51, 51)     # Dark Charcoal

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_title = p_title.add_run("LAPORAN PENGUJIAN OTOMATIS SISTEM TVJP")
    r_title.bold = True
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(22)
    r_title.font.color.rgb = COLOR_PRIMARY

    # Subtitle & Metadata
    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("Sistem Virtual Tutor Bahasa Jepang Berbasis Knowledge Graph\n")
    r_sub.italic = True
    r_sub.font.name = "Calibri"
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = RGBColor(100, 100, 100)

    r_meta1 = p_sub.add_run(f"Waktu Eksekusi: ")
    r_meta1.bold = True
    r_meta1.font.size = Pt(10)
    r_meta2 = p_sub.add_run(f"{now_str}   |   Status Suite: ")
    r_meta2.font.size = Pt(10)
    r_meta3 = p_sub.add_run(f"{status_text}\n")
    r_meta3.bold = True
    r_meta3.font.size = Pt(10)
    r_meta3.font.color.rgb = COLOR_SUCCESS if exitstatus == 0 else COLOR_FAIL

    doc.add_paragraph()

    # Executive Summary Heading
    p_h1 = doc.add_paragraph()
    r_h1 = p_h1.add_run("1. Executive Summary (Ringkasan Eksekutif)")
    r_h1.bold = True
    r_h1.font.name = "Calibri"
    r_h1.font.size = Pt(14)
    r_h1.font.color.rgb = COLOR_PRIMARY

    # Summary Table
    headers_sum = ["Layer", "Nama Pengujian", "Metodologi / Pilar", "Total", "Passed", "Failed", "Error", "Pass Rate", "Status"]
    table_sum = doc.add_table(rows=1, cols=len(headers_sum))
    table_sum.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_horizontal_borders_only(table_sum, border_color="1B365D", inside_color="D3D3D3")
    _set_table_cell_margins(table_sum, top=100, bottom=100, left=120, right=120)

    # Header Row Formatting
    hdr_cells = table_sum.rows[0].cells
    for idx, name in enumerate(headers_sum):
        _set_cell_background(hdr_cells[idx], "1B365D")
        align = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 3, 4, 5, 6, 7, 8) else WD_ALIGN_PARAGRAPH.LEFT
        _format_word_cell(hdr_cells[idx], name, bold=True, color=RGBColor(255, 255, 255), font_size=9.5, align=align)

    total_all, passed_all, failed_all, error_all = 0, 0, 0, 0

    for row_i, (l_id, l_name, l_desc, data, fmt) in enumerate(layers_data):
        t = len(data)
        p = sum(1 for r in data if r.get("status") == "PASSED")
        f = sum(1 for r in data if r.get("status") == "FAILED")
        e = sum(1 for r in data if r.get("status") == "ERROR")
        rate = (p / t * 100) if t > 0 else 0.0
        st_txt = "PASSED" if (f + e) == 0 and t > 0 else ("FAILED" if t > 0 else "NO DATA")

        total_all += t
        passed_all += p
        failed_all += f
        error_all += e

        row_cells = table_sum.add_row().cells
        bg = "F4F7FA" if row_i % 2 == 1 else "FFFFFF"
        for c in row_cells: _set_cell_background(c, bg)

        _format_word_cell(row_cells[0], l_id, bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _format_word_cell(row_cells[1], l_name, bold=False, font_size=9)
        _format_word_cell(row_cells[2], l_desc, italic=True, font_size=8.5, color=RGBColor(100,100,100))
        _format_word_cell(row_cells[3], str(t), font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _format_word_cell(row_cells[4], str(p), font_size=9, color=COLOR_SUCCESS, align=WD_ALIGN_PARAGRAPH.CENTER)
        _format_word_cell(row_cells[5], str(f), font_size=9, color=COLOR_FAIL if f > 0 else None, align=WD_ALIGN_PARAGRAPH.CENTER)
        _format_word_cell(row_cells[6], str(e), font_size=9, color=COLOR_FAIL if e > 0 else None, align=WD_ALIGN_PARAGRAPH.CENTER)
        _format_word_cell(row_cells[7], f"{rate:.1f}%", bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _format_word_cell(row_cells[8], st_txt, bold=True, color=COLOR_SUCCESS if st_txt=="PASSED" else COLOR_FAIL, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Total Summary Row
    overall_rate = (passed_all / total_all * 100) if total_all > 0 else 0.0
    overall_st = "SYSTEM VALID" if (failed_all + error_all) == 0 and total_all > 0 else "SYSTEM CACAT"
    tot_cells = table_sum.add_row().cells
    for c in tot_cells: _set_cell_background(c, "E6ECF5")

    _format_word_cell(tot_cells[0], "TOTAL", bold=True, font_size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    _format_word_cell(tot_cells[1], "Seluruh Layer (1-9)", bold=True, font_size=9.5)
    _format_word_cell(tot_cells[2], "Pengujian Komprehensif", italic=True, font_size=9)
    _format_word_cell(tot_cells[3], str(total_all), bold=True, font_size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    _format_word_cell(tot_cells[4], str(passed_all), bold=True, color=COLOR_SUCCESS, font_size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    _format_word_cell(tot_cells[5], str(failed_all), bold=True, color=COLOR_FAIL if failed_all > 0 else None, font_size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    _format_word_cell(tot_cells[6], str(error_all), bold=True, color=COLOR_FAIL if error_all > 0 else None, font_size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    _format_word_cell(tot_cells[7], f"{overall_rate:.1f}%", bold=True, font_size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    _format_word_cell(tot_cells[8], overall_st, bold=True, color=COLOR_SUCCESS if overall_st=="SYSTEM VALID" else COLOR_FAIL, font_size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # Detail Section Heading
    p_h2 = doc.add_paragraph()
    r_h2 = p_h2.add_run("2. Detail Hasil Pengujian per Layer")
    r_h2.bold = True
    r_h2.font.name = "Calibri"
    r_h2.font.size = Pt(14)
    r_h2.font.color.rgb = COLOR_PRIMARY

    for l_id, l_name, l_desc, data, fmt in layers_data:
        p_l = doc.add_paragraph()
        p_l.paragraph_format.space_before = Pt(12)
        p_l.paragraph_format.space_after = Pt(2)
        r_li = p_l.add_run(f"{l_id}: {l_name}")
        r_li.bold = True
        r_li.font.name = "Calibri"
        r_li.font.size = Pt(11)
        r_li.font.color.rgb = COLOR_PRIMARY

        p_ld = doc.add_paragraph()
        p_ld.paragraph_format.space_after = Pt(6)
        r_ld = p_ld.add_run(f"Fokus Pengujian: {l_desc}")
        r_ld.italic = True
        r_ld.font.size = Pt(9.5)
        r_ld.font.color.rgb = RGBColor(120, 120, 120)

        if not data:
            p_empty = doc.add_paragraph()
            r_emp = p_empty.add_run("Tidak ada data pengujian yang tercatat untuk layer ini.")
            r_emp.italic = True
            r_emp.font.size = Pt(9)
            continue

        if fmt == "standard":
            hdrs = ["No", "Service / Komponen", "Skenario / Kasus Uji", "Input Testing", "Ekspektasi Output", "Aktual Output", "Durasi", "Status"]
            tbl = doc.add_table(rows=1, cols=len(hdrs))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            _set_table_horizontal_borders_only(tbl, border_color="2B579A", inside_color="E0E0E0")
            _set_table_cell_margins(tbl, top=80, bottom=80, left=100, right=100)

            for idx, h in enumerate(hdrs):
                _set_cell_background(tbl.rows[0].cells[idx], "2B579A")
                align = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 6, 7) else WD_ALIGN_PARAGRAPH.LEFT
                _format_word_cell(tbl.rows[0].cells[idx], h, bold=True, color=RGBColor(255,255,255), font_size=8.5, align=align)

            for r_idx, item in enumerate(data):
                row_c = tbl.add_row().cells
                bg = "F9FBFD" if r_idx % 2 == 1 else "FFFFFF"
                for c in row_c: _set_cell_background(c, bg)

                st_val = item.get("status", "PASSED")
                _format_word_cell(row_c[0], str(item.get("no", "-")), font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
                _format_word_cell(row_c[1], item.get("service", "-"), bold=True, font_size=8)
                _format_word_cell(row_c[2], item.get("name", "-"), font_size=8)
                _format_word_cell(row_c[3], item.get("inputs", "-"), font_size=7.5, color=RGBColor(80,80,80))
                _format_word_cell(row_c[4], item.get("expected", "-"), font_size=7.5, color=RGBColor(80,80,80))
                _format_word_cell(row_c[5], item.get("actual", "-"), font_size=7.5, color=RGBColor(80,80,80))
                _format_word_cell(row_c[6], f"{item.get('duration', 0.0):.4f}s", font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
                _format_word_cell(row_c[7], st_val, bold=True, color=COLOR_SUCCESS if st_val=="PASSED" else COLOR_FAIL, font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

        elif fmt == "pbt":
            hdrs = ["No", "ID Properti", "Nama Properti Invarian", "Sampel Fuzzing", "Catatan Hasil / Invariant Check", "Status"]
            tbl = doc.add_table(rows=1, cols=len(hdrs))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            _set_table_horizontal_borders_only(tbl, border_color="2B579A", inside_color="E0E0E0")
            _set_table_cell_margins(tbl, top=80, bottom=80, left=100, right=100)

            for idx, h in enumerate(hdrs):
                _set_cell_background(tbl.rows[0].cells[idx], "2B579A")
                align = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 1, 3, 5) else WD_ALIGN_PARAGRAPH.LEFT
                _format_word_cell(tbl.rows[0].cells[idx], h, bold=True, color=RGBColor(255,255,255), font_size=8.5, align=align)

            for r_idx, item in enumerate(data):
                row_c = tbl.add_row().cells
                bg = "F9FBFD" if r_idx % 2 == 1 else "FFFFFF"
                for c in row_c: _set_cell_background(c, bg)

                st_val = item.get("status", "PASSED")
                _format_word_cell(row_c[0], str(item.get("no", "-")), font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
                _format_word_cell(row_c[1], item.get("id", "-"), bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
                _format_word_cell(row_c[2], item.get("name", "-"), font_size=8)
                _format_word_cell(row_c[3], f"~{item.get('examples', 0)} cases", font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
                _format_word_cell(row_c[4], item.get("note", "-"), font_size=7.5, color=RGBColor(80,80,80))
                _format_word_cell(row_c[5], st_val, bold=True, color=COLOR_SUCCESS if st_val=="PASSED" else COLOR_FAIL, font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

        elif fmt == "mt":
            hdrs = ["No", "ID Relasi", "Nama Metamorphic Relation", "Kasus Uji", "Transformasi Input (T)", "Relasi Output (R)", "Status"]
            tbl = doc.add_table(rows=1, cols=len(hdrs))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            _set_table_horizontal_borders_only(tbl, border_color="2B579A", inside_color="E0E0E0")
            _set_table_cell_margins(tbl, top=80, bottom=80, left=100, right=100)

            for idx, h in enumerate(hdrs):
                _set_cell_background(tbl.rows[0].cells[idx], "2B579A")
                align = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 1, 3, 6) else WD_ALIGN_PARAGRAPH.LEFT
                _format_word_cell(tbl.rows[0].cells[idx], h, bold=True, color=RGBColor(255,255,255), font_size=8.5, align=align)

            for r_idx, item in enumerate(data):
                row_c = tbl.add_row().cells
                bg = "F9FBFD" if r_idx % 2 == 1 else "FFFFFF"
                for c in row_c: _set_cell_background(c, bg)

                st_val = item.get("status", "PASSED")
                _format_word_cell(row_c[0], str(item.get("no", "-")), font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
                _format_word_cell(row_c[1], item.get("id", "-"), bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
                _format_word_cell(row_c[2], item.get("name", "-"), font_size=8)
                _format_word_cell(row_c[3], f"{item.get('cases', 0)} cases", font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
                _format_word_cell(row_c[4], item.get("transform", "-"), font_size=7.5, color=RGBColor(80,80,80))
                _format_word_cell(row_c[5], item.get("relation", "-"), font_size=7.5, color=RGBColor(80,80,80))
                _format_word_cell(row_c[6], st_val, bold=True, color=COLOR_SUCCESS if st_val=="PASSED" else COLOR_FAIL, font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

        doc.add_paragraph()

    try:
        doc.save(docx_file)
    except PermissionError:
        print("\n[WARNING] File 'laporan.docx' sedang dibuka di Microsoft Word. Menyimpan ke 'laporan_latest.docx'...")
        alt_docx = os.path.join(report_dir, "Pengujian.docx")
        try:
            doc.save(alt_docx)
        except Exception:
            pass

    # ─────────────────────────────────────────────────────────────────────────
    # 2. GENERATE MARKDOWN DOCUMENT (.md) FOR DUAL SUPPORT
    # ─────────────────────────────────────────────────────────────────────────
    lines = []
    lines.append("# 📊 LAPORAN PENGUJIAN OTOMATIS SISTEM TVJP")
    lines.append("")
    lines.append("> **Sistem Virtual Tutor Bahasa Jepang Berbasis Knowledge Graph**  ")
    lines.append(f"> **Waktu Eksekusi Pengujian:** `{now_str}`  ")
    lines.append(f"> **Status Suite:** `{status_text}`")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## 📈 Executive Summary (Ringkasan Eksekutif)")
    lines.append("")
    lines.append("| Layer | Nama Pengujian | Metodologi / Pilar | Total Tes | Passed | Failed | Error | Pass Rate | Status |")
    lines.append("| :---: | :--- | :--- | :---: | :---: | :---: | :---: | :---: | :---: |")

    for l_id, l_name, l_desc, data, fmt in layers_data:
        t = len(data)
        p = sum(1 for r in data if r.get("status") == "PASSED")
        f = sum(1 for r in data if r.get("status") == "FAILED")
        e = sum(1 for r in data if r.get("status") == "ERROR")
        rate = (p / t * 100) if t > 0 else 0.0
        st_icon = "✅ PASSED" if (f + e) == 0 and t > 0 else ("❌ FAILED" if t > 0 else "⚠️ NO DATA")
        lines.append(f"| **{l_id}** | {l_name} | {l_desc} | {t} | {p} | {f} | {e} | {rate:.1f}% | {st_icon} |")

    lines.append(f"| **TOTAL** | **Seluruh Layer (1-9)** | **Pengujian Komprehensif** | **{total_all}** | **{passed_all}** | **{failed_all}** | **{error_all}** | **{overall_rate:.1f}%** | **{overall_st}** |")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## 📝 Detail Hasil Pengujian per Layer")
    lines.append("")

    for l_id, l_name, l_desc, data, fmt in layers_data:
        lines.append(f"### 🔹 {l_id}: {l_name}")
        lines.append(f"*Fokus Pengujian: {l_desc}*")
        lines.append("")

        if not data:
            lines.append("> *Tidak ada data pengujian yang tercatat untuk layer ini.*")
            lines.append("")
            continue

        if fmt == "standard":
            lines.append("| No | Service / Komponen | Skenario / Kasus Uji | Input Testing | Ekspektasi Output | Aktual Output | Durasi | Status |")
            lines.append("| :---: | :--- | :--- | :--- | :--- | :--- | :---: | :---: |")
            for item in data:
                st_str = "✅ PASSED" if item.get("status") == "PASSED" else "❌ FAIL"
                lines.append(f"| {item.get('no', '-')} | **{item.get('service', '-')}** | {item.get('name', '-')} | `{item.get('inputs', '-')}` | `{item.get('expected', '-')}` | `{item.get('actual', '-')}` | {item.get('duration', 0.0):.4f}s | {st_str} |")
        elif fmt == "pbt":
            lines.append("| No | ID Properti | Nama Properti Invarian | Sampel Fuzzing | Catatan Hasil / Invariant Check | Status |")
            lines.append("| :---: | :---: | :--- | :---: | :--- | :---: |")
            for item in data:
                st_str = "✅ PASSED" if item.get("status") == "PASSED" else "❌ FAIL"
                lines.append(f"| {item.get('no', '-')} | **{item.get('id', '-')}** | {item.get('name', '-')} | ~{item.get('examples', 0)} cases | {item.get('note', '-')} | {st_str} |")
        elif fmt == "mt":
            lines.append("| No | ID Relasi | Nama Metamorphic Relation | Kasus Uji | Transformasi Input (T) | Relasi Output (R) | Status |")
            lines.append("| :---: | :---: | :--- | :---: | :--- | :--- | :---: |")
            for item in data:
                st_str = "✅ PASSED" if item.get("status") == "PASSED" else "❌ FAIL"
                lines.append(f"| {item.get('no', '-')} | **{item.get('id', '-')}** | {item.get('name', '-')} | {item.get('cases', 0)} cases | `{item.get('transform', '-')}` | `{item.get('relation', '-')}` | {st_str} |")
        lines.append("")

    with open(md_file, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
