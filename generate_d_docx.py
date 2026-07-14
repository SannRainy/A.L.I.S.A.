# -*- coding: utf-8 -*-
"""
generate_d_docx.py
==================
Python script to generate "D. Pengujian" section of the TVJP report in Word format (.docx).
Uses custom academic styling with clean borders, light gray headers, and crisp margins.
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex: str):
    """Set cell background color (shading)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_table_horizontal_borders(table, border_color="7F7F7F", inside_color="D3D3D3"):
    """
    Sets table borders so only horizontal lines are visible (academic style).
    """
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="12" w:space="0" w:color="{border_color}"/>
        <w:bottom w:val="single" w:sz="12" w:space="0" w:color="{border_color}"/>
        <w:left w:val="none"/>
        <w:right w:val="none"/>
        <w:insideH w:val="single" w:sz="6" w:space="0" w:color="{inside_color}"/>
        <w:insideV w:val="none"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))

def set_table_cell_margins(table, top=120, bottom=120, left=140, right=140):
    """Sets cell padding for better readability."""
    tblPr = table._tbl.tblPr
    margins_xml = f'''
    <w:tblCellMar {nsdecls("w")}>
        <w:top w:w="{top}" w:type="dxa"/>
        <w:bottom w:w="{bottom}" w:type="dxa"/>
        <w:left w:w="{left}" w:type="dxa"/>
        <w:right w:w="{right}" w:type="dxa"/>
    </w:tblCellMar>
    '''
    tblPr.append(parse_xml(margins_xml))

def format_word_cell(cell, text: str, bold=False, italic=False, color=None, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Formats cell text consistently."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(str(text))
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color

def create_document():
    doc = docx.Document()
    
    # Page settings (1 inch margin)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Heading 1
    p_h1 = doc.add_paragraph()
    run_h1 = p_h1.add_run("D. Pengujian Sistem")
    run_h1.bold = True
    run_h1.font.name = "Times New Roman"
    run_h1.font.size = Pt(14)
    p_h1.paragraph_format.space_before = Pt(12)
    p_h1.paragraph_format.space_after = Pt(6)

    # Heading 2
    p_h2 = doc.add_paragraph()
    run_h2 = p_h2.add_run("1. Black-box testing")
    run_h2.bold = True
    run_h2.font.name = "Times New Roman"
    run_h2.font.size = Pt(12)
    p_h2.paragraph_format.space_before = Pt(12)
    p_h2.paragraph_format.space_after = Pt(6)

    # Paragraph intro
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.line_spacing = 1.5
    p_intro.paragraph_format.space_after = Pt(12)
    p_intro.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_intro = p_intro.add_run(
        "Pengujian Black-box (kotak hitam) digunakan untuk menguji fungsionalitas sistem TVJP A.L.I.S.A. "
        "tanpa harus mengetahui struktur internal kode program. Pengujian difokuskan pada input yang diberikan "
        "dan output yang dihasilkan untuk memastikan bahwa semua fitur berjalan sesuai dengan skenario perancangan "
        "antarmuka yang telah ditentukan sebelumnya."
    )
    run_intro.font.name = "Times New Roman"
    run_intro.font.size = Pt(11)

    # Testing Data for 11 interfaces
    testing_data = [
        {
            "num": "1.1",
            "page_name": "Testing Halaman Login",
            "description": "Pengujian halaman login dilakukan untuk memverifikasi proses autentikasi pengguna masuk ke dalam sistem dengan menggunakan kredensial berupa surel (email) dan kata sandi.",
            "tabel_title": "Tabel 4.1 Testing Halaman Login",
            "cases": [
                {
                    "no": "1",
                    "skenario": "Login sukses pengguna umum",
                    "case": "1. Input email pengguna\n2. Input password valid\n3. Klik tombol Sign in",
                    "data": "Email: dimasprayogo@student.com\nPassword: Password123 (password valid)",
                    "expected": "1. Sistem menampilkan pesan berhasil login\n2. Pengguna diarahkan ke halaman Homepage Chatbot",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "2",
                    "skenario": "Login sukses admin",
                    "case": "1. Input email admin\n2. Input password admin\n3. Klik tombol Sign in",
                    "data": "Email: admin@admin.com\nPassword: Password123",
                    "expected": "1. Sistem menampilkan pesan berhasil login\n2. Admin diarahkan ke halaman Dashboard Admin",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "3",
                    "skenario": "Login gagal karena email tidak terdaftar",
                    "case": "1. Input email\n2. Input password\n3. Klik tombol Sign in",
                    "data": "Email: salah@example.com\nPassword: test",
                    "expected": "1. Sistem menampilkan pesan validasi gagal\n2. Pengguna tetap berada pada halaman",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                }
            ]
        },
        {
            "num": "1.2",
            "page_name": "Testing Halaman Register",
            "description": "Pengujian halaman register dilakukan untuk memverifikasi kelancaran pembuatan akun baru dan ketepatan perekaman data profil demografis belajar.",
            "tabel_title": "Tabel 4.2 Testing Halaman Register",
            "cases": [
                {
                    "no": "1",
                    "skenario": "Register sukses dengan profil lengkap",
                    "case": "1. Isi formulir akun\n2. Isi umur, gender, asal negara, tujuan belajar\n3. Klik tombol Daftar",
                    "data": "Email: baru@gmail.com\nPass: Pass123\nNama: Rian\nUmur: 21, Gender: Laki-laki\nNegara: Indonesia\nTujuan: Hobi",
                    "expected": "1. Akun berhasil terbuat\n2. Notifikasi sukses muncul\n3. Pengguna dialihkan ke login overlay",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "2",
                    "skenario": "Register gagal karena email sudah terdaftar",
                    "case": "1. Masukkan email terdaftar\n2. Isi formulir lainnya\n3. Klik tombol Daftar",
                    "data": "Email: dimasprayogo@student.com\nPass: Pass123",
                    "expected": "1. Notifikasi error email sudah terdaftar tampil\n2. Proses pendaftaran dibatalkan",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "3",
                    "skenario": "Register gagal karena password kurang dari 6 karakter",
                    "case": "1. Isi email baru\n2. Isi password pendek\n3. Klik tombol Daftar",
                    "data": "Email: pendek@gmail.com\nPass: abc",
                    "expected": "1. Validasi error password minimal 6 karakter tampil\n2. Pengguna tetap berada di halaman",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                }
            ]
        },
        {
            "num": "1.3",
            "page_name": "Testing Halaman Dashboard Admin",
            "description": "Pengujian halaman dashboard admin dilakukan untuk menguji pembatasan hak akses berbasis peran (RBAC) serta pemuatan grafik statistik analitik.",
            "tabel_title": "Tabel 4.3 Testing Halaman Dashboard Admin",
            "cases": [
                {
                    "no": "1",
                    "skenario": "Akses dashboard admin dengan akun administrator",
                    "case": "1. Login dengan akun role admin\n2. Arahkan rute ke /admin",
                    "data": "Email: admin@admin.com\nRole: admin",
                    "expected": "1. Sistem memvalidasi role admin\n2. Halaman admin shell dimuat sepenuhnya dengan seluruh tab data",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "2",
                    "skenario": "Akses dashboard admin ditolak untuk pengguna biasa",
                    "case": "1. Login dengan akun role user\n2. Arahkan rute ke /admin secara manual",
                    "data": "Email: dimasprayogo@student.com\nRole: user",
                    "expected": "1. Sistem menolak akses masuk\n2. Pengguna dialihkan kembali ke homepage utama secara otomatis",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "3",
                    "skenario": "Memuat data analitik di tab grafik",
                    "case": "1. Klik tab Analisis & Grafik\n2. Tunggu proses loading grafik selesai",
                    "data": "Membuka menu grafik ringkasan pengguna dan AI models",
                    "expected": "1. Sistem merender diagram lingkaran sebaran negara dan diagram batang skor kuis secara akurat",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                }
            ]
        },
        {
            "num": "1.4",
            "page_name": "Testing Halaman CSV File Management",
            "description": "Pengujian halaman CSV File Management dilakukan untuk memverifikasi proses pengelolaan kurikulum serta sinkronisasi data ke database graf Neo4j.",
            "tabel_title": "Tabel 4.4 Testing Halaman CSV File Management",
            "cases": [
                {
                    "no": "1",
                    "skenario": "Memuat daftar berkas CSV kurikulum",
                    "case": "1. Masuk halaman Admin\n2. Klik tab Data Pipeline",
                    "data": "Permintaan pemuatan berkas dari backend API",
                    "expected": "1. Sistem menampilkan daftar file CSV kurikulum (kanji.csv, vocab.csv, grammar.csv, topic.csv)",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "2",
                    "skenario": "Mengedit dan menyimpan data kosakata langsung pada tabel",
                    "case": "1. Klik berkas vocab.csv\n2. Ubah salah satu sel kosakata\n3. Klik tombol Save CSV",
                    "data": "Mengubah arti kata '食べる' menjadi 'Makan (bentuk kamus)'",
                    "expected": "1. Berkas vocab.csv berhasil diperbarui di server\n2. Pesan sukses ditampilkan di layar",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "3",
                    "skenario": "Menjalankan pipeline sinkronisasi Neo4j",
                    "case": "1. Klik tombol Trigger Ingest\n2. Tunggu proses integrasi backend selesai",
                    "data": "Mengirim permintaan POST ke API /ingest",
                    "expected": "1. Sistem melakukan sinkronisasi data CSV ke Neo4j\n2. Menampilkan status sukses integrasi simpul dan relasi graf",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                }
            ]
        },
        {
            "num": "1.5",
            "page_name": "Testing Homepage Chatbot Mode",
            "description": "Pengujian chatbot mode dilakukan untuk memastikan obrolan teks, sinkronisasi model avatar 3D, dan suggestion chips berfungsi dengan baik.",
            "tabel_title": "Tabel 4.5 Testing Homepage Chatbot Mode",
            "cases": [
                {
                    "no": "1",
                    "skenario": "Mengirim pesan teks ke tutor virtual",
                    "case": "1. Tulis pesan di kotak input\n2. Klik tombol Send atau tekan Enter",
                    "data": "Pesan: 'Tolong jelaskan cara memakai partikel に'",
                    "expected": "1. Teks pesan tampil di area percakapan\n2. Tutor virtual memberikan respon balasan berupa penjelasan materi",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "2",
                    "skenario": "Streaming tanggapan teks dan visualisasi avatar 3D",
                    "case": "1. Amati tampilan saat tutor merespons\n2. Perhatikan pergerakan mulut avatar",
                    "data": "Penerimaan teks bertahap kata demi kata (streaming SSE)",
                    "expected": "1. Teks respons tampil secara mengalir\n2. Model 3D VRM menggerakkan bibir (lip-sync) beriringan dengan audio",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "3",
                    "skenario": "Menggunakan Suggestion Chips topik obrolan",
                    "case": "1. Amati chip rekomendasi di bawah input\n2. Klik salah satu chip yang tersedia",
                    "data": "Membuka chip: 'Bisa berikan contoh kalimat?'",
                    "expected": "1. Teks chip otomatis terkirim sebagai pesan baru\n2. Tutor merespons dengan contoh kalimat bahasa Jepang yang tepat",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                }
            ]
        },
        {
            "num": "1.6",
            "page_name": "Testing Homepage Quiz Mode",
            "description": "Pengujian quiz mode dilakukan untuk memverifikasi perenderan peta level belajar sekuensial, pengerjaan kuis, dan pembaruan tingkat kemahiran kognitif BKT.",
            "tabel_title": "Tabel 4.6 Testing Homepage Quiz Mode",
            "cases": [
                {
                    "no": "1",
                    "skenario": "Membuka peta level kuis sekuensial",
                    "case": "1. Pilih tab Quest Mode\n2. Perhatikan status simpul pada peta belajar",
                    "data": "Data progres belajar pengguna terkini",
                    "expected": "1. Peta dirender dalam bentuk rangkaian simpul materi\n2. Simpul materi yang belum memenuhi prasyarat terkunci otomatis",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "2",
                    "skenario": "Mengerjakan kuis latihan pilihan ganda",
                    "case": "1. Klik simpul level yang terbuka\n2. Pilih salah satu jawaban\n3. Klik tombol Submit",
                    "data": "Pengerjaan latihan kosakata level quest_n5_vocab_1",
                    "expected": "1. Sistem menampilkan feedback benar/salah secara instan\n2. Bilah kemajuan bertambah",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "3",
                    "skenario": "Menyimpan hasil kuis dan memperbarui kemahiran BKT",
                    "case": "1. Jawab semua soal kuis\n2. Amati halaman hasil kuis",
                    "data": "Menyelesaikan 5 soal kuis dengan skor akhir 80",
                    "expected": "1. Hasil kuis tersimpan di database Supabase\n2. Backend memperbarui nilai probabilitas penguasaan BKT siswa",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                }
            ]
        },
        {
            "num": "1.7",
            "page_name": "Testing Halaman Kanji Dojo",
            "description": "Pengujian Kanji Dojo dilakukan untuk memverifikasi fungsionalitas pemutaran audio pelafalan kanji, pembalikan flashcard, dan animasi urutan goresan kanji.",
            "tabel_title": "Tabel 4.7 Testing Halaman Kanji Dojo",
            "cases": [
                {
                    "no": "1",
                    "skenario": "Mempelajari cara baca dan arti lewat flashcard kanji",
                    "case": "1. Masuk menu Kanji Dojo\n2. Klik flashcard kanji untuk membalikkan kartu",
                    "data": "Membuka kartu kanji '水' (Air)",
                    "expected": "1. Flashcard berputar 3D menampilkan cara baca onyomi/kunyomi dan terjemahan bahasa Indonesia di sisi belakang",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "2",
                    "skenario": "Memutar suara pelafalan kanji",
                    "case": "1. Klik tombol ikon speaker suara pada flashcard kanji",
                    "data": "Teks kanji yang akan dilafalkan",
                    "expected": "1. Audio pelafalan kanji Jepang diputar dengan jelas",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "3",
                    "skenario": "Menampilkan animasi cara penulisan kanji",
                    "case": "1. Amati panel animasi goresan kanji\n2. Tekan tombol Play Stroke Animation",
                    "data": "File SVG animasi goresan kanji terpilih",
                    "expected": "1. Sistem menggambarkan urutan goresan pena penulisan kanji langkah demi langkah",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                }
            ]
        },
        {
            "num": "1.8",
            "page_name": "Testing Halaman SRS Review",
            "description": "Pengujian SRS review dilakukan untuk menguji keandalan antrean kartu hafalan serta kalkulasi tanggal ulasan berikutnya berdasarkan algoritma SM-2.",
            "tabel_title": "Tabel 4.8 Testing Halaman SRS Review",
            "cases": [
                {
                    "no": "1",
                    "skenario": "Memuat kartu materi ulasan jeda",
                    "case": "1. Buka menu SRS Review",
                    "data": "Daftar kosakata/tata bahasa terjadwal ulasan hari ini",
                    "expected": "1. Kartu ulasan materi tampil di layar\n2. Jumlah antrean ulasan terhitung dengan benar",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "2",
                    "skenario": "Menilai pemahaman diri dengan tombol rating SM-2",
                    "case": "1. Klik Tampilkan Jawaban\n2. Klik salah satu tombol rating (misal: 'Good')",
                    "data": "Menilai pemahaman kosakata '食べる' dengan tombol 'Good'",
                    "expected": "1. Data repetisi diperbarui\n2. Tanggal ulasan berikutnya dihitung ulang secara otomatis\n3. Kartu berganti ke antrean berikutnya",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "3",
                    "skenario": "Penjadwalan ulasan ulang ketika materi dinilai belum paham",
                    "case": "1. Klik Tampilkan Jawaban\n2. Klik tombol 'Again'",
                    "data": "Memilih tombol 'Again' (belum paham)",
                    "expected": "1. Repetisi direset ke awal\n2. Kartu materi tersebut dimasukkan kembali ke antrean ulasan terdekat pada hari yang sama",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                }
            ]
        },
        {
            "num": "1.9",
            "page_name": "Testing Homepage Speaking Mode",
            "description": "Pengujian speaking mode dilakukan untuk menguji keakuratan transkripsi ucapan (STT) serta modul penilaian akurasi pelafalan dari backend.",
            "tabel_title": "Tabel 4.9 Testing Homepage Speaking Mode",
            "cases": [
                {
                    "no": "1",
                    "skenario": "Mengaktifkan mikrofon dan merekam suara siswa",
                    "case": "1. Masuk menu Voice Mode\n2. Klik dan tahan tombol mic untuk berbicara",
                    "data": "Rekaman audio input dari mic",
                    "expected": "1. Cincin visualisasi gelombang suara aktif bergerak\n2. Perekaman berjalan tanpa eror visual di UI",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "2",
                    "skenario": "Pengenalan suara secara real-time (Speech-to-Text)",
                    "case": "1. Ucapkan frasa bahasa Jepang ke mikrofon\n2. Lepaskan tombol perekaman",
                    "data": "Ucapan siswa: 'こんにちは'",
                    "expected": "1. Sistem mentranskripsikan suara menjadi teks Jepang secara instan di layar",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "3",
                    "skenario": "Penilaian akurasi pengucapan lisan",
                    "case": "1. Ucapkan kalimat latihan lisan\n2. Tunggu hasil penilaian dari backend",
                    "data": "Latihan kalimat: 'これはペンです'",
                    "expected": "1. Sistem menampilkan persentase akurasi pengucapan siswa (misal: 92% akurasi)\n2. Evaluasi detail pelafalan ditampilkan",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                }
            ]
        },
        {
            "num": "1.10",
            "page_name": "Testing Halaman Profile",
            "description": "Pengujian halaman profile dilakukan untuk memverifikasi pemuatan identitas pengguna, visualisasi diagram radar, dan pengeksporan data belajar.",
            "tabel_title": "Tabel 4.10 Testing Halaman Profile",
            "cases": [
                {
                    "no": "1",
                    "skenario": "Memuat visualisasi radar chart kompetensi belajar",
                    "case": "1. Masuk ke halaman Profile\n2. Amati grafik kompetensi",
                    "data": "Nilai kemahiran BKT pada 5 aspek bahasa (Kanji, Vocab, Grammar, Listening, Speaking)",
                    "expected": "1. Diagram radar 5 dimensi dirender secara dinamis berdasarkan data kemahiran aktual",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "2",
                    "skenario": "Menampilkan streak kalender dan statistik harian",
                    "case": "1. Buka halaman Profile\n2. Perhatikan streak harian",
                    "data": "Riwayat aktivitas belajar di Supabase",
                    "expected": "1. Kalender keaktifan memunculkan tanda centang pada tanggal aktif belajar\n2. Jumlah streak beruntun dihitung secara akurat",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "3",
                    "skenario": "Ekspor data riwayat belajar pengguna",
                    "case": "1. Klik tombol Ekspor Data",
                    "data": "Pengunduhan file data profil terstruktur",
                    "expected": "1. Berkas JSON berisi riwayat progres belajar lengkap berhasil diunduh",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                }
            ]
        },
        {
            "num": "1.11",
            "page_name": "Testing Halaman Achievement",
            "description": "Pengujian halaman achievement dilakukan untuk memastikan tampilan gamifikasi lencana dan kenaikan pangkat kepangkatan samurai berfungsi dengan benar.",
            "tabel_title": "Tabel 4.11 Testing Halaman Achievement",
            "cases": [
                {
                    "no": "1",
                    "skenario": "Menampilkan lencana penghargaan (badges) yang diperoleh",
                    "case": "1. Masuk ke halaman Achievement\n2. Amati daftar lencana",
                    "data": "Data pencapaian lencana siswa (user_badges)",
                    "expected": "1. Lencana yang telah diperoleh tampil berwarna cerah\n2. Lencana yang belum didapatkan tampil redup (grayscale)",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "2",
                    "skenario": "Memasang emblem pilihan (equipped emblems)",
                    "case": "1. Pilih lencana yang telah diraih\n2. Klik opsi pasang emblem",
                    "data": "Memasang lencana 'Kanji Dojo Master'",
                    "expected": "1. Lencana terpajang di area emblem profil terpasang\n2. Perubahan tersimpan di database",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                },
                {
                    "no": "3",
                    "skenario": "Kenaikan pangkat samurai (samurai ranks) berdasarkan XP",
                    "case": "1. Kumpulkan poin XP tambahan dari kuis\n2. Periksa perubahan gelar kepangkatan di halaman Achievement",
                    "data": "Akumulasi XP bertambah melewati batas threshold pangkat berikutnya",
                    "expected": "1. Gelar pangkat samurai pengguna naik (misalnya dari Ashigaru ke Samurai)\n2. Visualisasi status pangkat diperbarui",
                    "result": "Berhasil sesuai harapan",
                    "conclusion": "[√] Valid"
                }
            ]
        }
    ]

    for p_idx, page in enumerate(testing_data):
        # Heading 3 (1.X Testing Halaman XXX)
        p_h3 = doc.add_paragraph()
        run_h3 = p_h3.add_run(f"{page['num']}. {page['page_name']}")
        run_h3.bold = True
        run_h3.font.name = "Times New Roman"
        run_h3.font.size = Pt(11)
        p_h3.paragraph_format.space_before = Pt(8)
        p_h3.paragraph_format.space_after = Pt(4)

        # Brief description
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.line_spacing = 1.15
        p_desc.paragraph_format.space_after = Pt(6)
        p_desc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_desc = p_desc.add_run(page['description'])
        run_desc.font.name = "Times New Roman"
        run_desc.font.size = Pt(10)

        # Add Table Title/Caption
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(4)
        p_cap.paragraph_format.space_after = Pt(2)
        run_cap = p_cap.add_run(page['tabel_title'])
        run_cap.italic = True
        run_cap.font.name = "Times New Roman"
        run_cap.font.size = Pt(10)

        # Create Table (1 header + len(cases) rows, 7 columns)
        table = doc.add_table(rows=1 + len(page['cases']), cols=7)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set borders & cell margins
        set_table_horizontal_borders(table, border_color="000000", inside_color="BFBFBF")
        set_table_cell_margins(table, top=100, bottom=100, left=120, right=120)

        # Column Headers
        headers = ["No", "Skenario", "Case", "Data", "Hasil yang Diharapkan", "Hasil Pengujian", "Kesimpulan"]
        col_widths = [Inches(0.4), Inches(1.2), Inches(1.3), Inches(1.3), Inches(1.4), Inches(0.9), Inches(0.8)]
        
        hdr_cells = table.rows[0].cells
        for i, text in enumerate(headers):
            hdr_cells[i].width = col_widths[i]
            # Light gray background (AEAAAA / D3D3D3) for header
            set_cell_background(hdr_cells[i], "D3D3D3")
            format_word_cell(hdr_cells[i], text, bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Populate rows
        for r_idx, case in enumerate(page['cases']):
            row_cells = table.rows[r_idx + 1].cells
            
            # Align width
            for i in range(7):
                row_cells[i].width = col_widths[i]

            # Populate content
            format_word_cell(row_cells[0], case['no'], font_size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
            format_word_cell(row_cells[1], case['skenario'], font_size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT)
            format_word_cell(row_cells[2], case['case'], font_size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT)
            format_word_cell(row_cells[3], case['data'], font_size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT)
            format_word_cell(row_cells[4], case['expected'], font_size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT)
            format_word_cell(row_cells[5], case['result'], font_size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT)
            format_word_cell(row_cells[6], case['conclusion'], font_size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Spacing after table
        doc.add_paragraph().paragraph_format.space_before = Pt(6)

    # Save document
    output_fn = "D_Pengujian_TVJP.docx"
    doc.save(output_fn)
    print(f"Success: Generated {output_fn}")

if __name__ == "__main__":
    create_document()
