import os
import shutil
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Define source and destination paths
src_path = r"SKRIPSI 2026\BAB 5 2205101053.docx"
dest_path = r"SKRIPSI 2026\BAB V PENUTUP.docx"

# Copy the file
shutil.copy(src_path, dest_path)
print(f"Copied {src_path} to {dest_path}")

doc = docx.Document(dest_path)

def format_text_run(run, font_name="Times New Roman", font_size=12, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic

# We modify Paragraph 3 (Kesimpulan) -> A. Kesimpulan
p3 = doc.paragraphs[3]
p3.text = "A. Kesimpulan"
p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
p3.paragraph_format.line_spacing = 2.0
p3.paragraph_format.left_indent = Inches(0.197222)
p3.paragraph_format.first_line_indent = None
# Reapply runs formatting
for run in p3.runs:
    format_text_run(run, bold=True)

# Now, we want to insert the intro and list items for Kesimpulan before Paragraph 4 (Saran)
p4 = doc.paragraphs[4]

# Under A. Kesimpulan:
intro_kesimpulan = "Berdasarkan Hasil Penelitian Sistem Tutor Pembelajaran Bahasa Jepang Berbasis Neuro-Symbolic AI Berikut :"
items_kesimpulan = [
    "1.\tSistem Tutor Pembelajaran Bahasa Jepang Berbasis Neuro-Symbolic AI (yang diwujudkan dalam aplikasi A.L.I.S.A.) berhasil dirancang dan diimplementasikan secara hibrida dengan memadukan kekuatan generatif model bahasa besar (Large Language Model) lokal Qwen untuk interaksi natural dan kepastian informasi dari basis pengetahuan terstruktur Knowledge Graph berbasis Neo4j melalui metode GraphRAG (Graph Retrieval-Augmented Generation) untuk menghindari halusinasi informasi.",
    "2.\tPenyusunan jalur pembelajaran terstruktur (Structured Learning Path) pada sistem ini berhasil diotomatisasi secara dinamis menggunakan algoritma topological sort pada Knowledge Graph berdasarkan dependensi prasyarat materi. Evaluasi tingkat pemahaman pembelajar juga berhasil diukur secara berkala dan real-time menggunakan algoritma Bayesian Knowledge Tracing (BKT), serta retensi memori jangka panjang pembelajar dioptimalkan secara adaptif melalui Spaced Repetition System (SRS) dengan algoritma SM-2.",
    "3.\tAntarmuka asisten virtual pada sistem ini berhasil dikembangkan secara interaktif dengan mengintegrasikan visualisasi graf pengetahuan interaktif, visualisasi avatar humanoid 3D VRM yang mampu menampilkan ekspresi tutor secara real-time, serta fitur mode percakapan vokal yang didukung oleh teknologi Text-to-Speech (TTS) dan Speech-to-Text (STT) untuk memfasilitasi latihan dialog bahasa Jepang terapan.",
    "4.\tPengujian sistem secara menyeluruh dengan metode Black-box, White-box, Property-Based Testing (PBT), dan Metamorphic Testing (MT) membuktikan bahwa sistem berfungsi dengan baik, stabil, dan memiliki ketahanan tinggi. Seluruh skenario pengujian fungsionalitas sukses 100%, logika unit testing terverifikasi dengan cakupan optimal, pengujian PBT membuktikan ketahanan terhadap masukan data ekstrem, dan pengujian MT berhasil memecahkan Oracle Problem pada komponen AI dengan membuktikan konsistensi model kognitif terhadap teori kognitif secara matematis."
]

# Insert intro paragraph before p4
p_intro_k = p4.insert_paragraph_before()
p_intro_k.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_intro_k.paragraph_format.line_spacing = 2.0
p_intro_k.paragraph_format.left_indent = Inches(0.197222)
p_intro_k.paragraph_format.first_line_indent = Inches(0.590972)
run_intro_k = p_intro_k.add_run(intro_kesimpulan)
format_text_run(run_intro_k)

# Insert list paragraphs before p4
for item in items_kesimpulan:
    p_item = p4.insert_paragraph_before()
    p_item.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_item.paragraph_format.line_spacing = 2.0
    p_item.paragraph_format.left_indent = Inches(0.39375)
    p_item.paragraph_format.first_line_indent = Inches(-0.196527)
    run_item = p_item.add_run(item)
    format_text_run(run_item)

# Insert an empty spacing paragraph before p4 to separate Kesimpulan list from Saran heading
p_space = p4.insert_paragraph_before()
p_space.paragraph_format.line_spacing = 2.0
p_space.paragraph_format.left_indent = Inches(0.197222)

# Modify Paragraph 4 (Saran) -> B. Saran
p4.text = "B. Saran"
p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
p4.paragraph_format.line_spacing = 2.0
p4.paragraph_format.left_indent = Inches(0.197222)
p4.paragraph_format.first_line_indent = None
for run in p4.runs:
    format_text_run(run, bold=True)

# Under B. Saran, since p4 is the last paragraph, we append new paragraphs at the end of the document
intro_saran = "Berdasarkan Hasil Penelitian Sistem Tutor Pembelajaran Bahasa Jepang Berbasis Neuro-Symbolic AI Berikut :"
items_saran = [
    "1.\tDisarankan untuk melakukan perluasan cakupan materi pembelajaran dan basis pengetahuan (Knowledge Graph) pada sistem ke tingkat yang lebih tinggi, seperti tingkat JLPT N4 hingga JLPT N1, agar sistem tutor virtual ini dapat dimanfaatkan secara berkelanjutan oleh pembelajar mandiri pada tingkat menengah dan lanjut.",
    "2.\tDisarankan untuk melakukan proses fine-tuning pada model bahasa besar (Large Language Model) lokal menggunakan dataset percakapan dan tata bahasa Jepang yang terspesialisasi, guna meningkatkan akurasi penjelasan tata bahasa Jepang serta meminimalkan latensi pemrosesan respon pada backend untuk interaksi yang lebih responsif.",
    "3.\tDisarankan untuk meningkatkan animasi visual dan sinkronisasi gerakan bibir (lip-sync) pada avatar humanoid 3D VRM agar ekspresi wajah tutor virtual dapat beradaptasi secara dinamis dan lebih presisi sesuai dengan intonasi suara yang disintesis.",
    "4.\tDisarankan untuk melaksanakan penelitian lapangan lanjutan berupa uji coba eksperimental secara langsung (field testing) kepada kelompok pengguna (mahasiswa atau pembelajar mandiri) dengan membandingkan kelompok kontrol dan kelompok eksperimen dalam kurun waktu tertentu, guna mengukur efektivitas Sistem Tutor Pembelajaran Bahasa Jepang Berbasis Neuro-Symbolic AI terhadap peningkatan hasil belajar secara empiris."
]

# Append intro paragraph
p_intro_s = doc.add_paragraph()
p_intro_s.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_intro_s.paragraph_format.line_spacing = 2.0
p_intro_s.paragraph_format.left_indent = Inches(0.197222)
p_intro_s.paragraph_format.first_line_indent = Inches(0.590972)
run_intro_s = p_intro_s.add_run(intro_saran)
format_text_run(run_intro_s)

# Append list items
for item in items_saran:
    p_item = doc.add_paragraph()
    p_item.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_item.paragraph_format.line_spacing = 2.0
    p_item.paragraph_format.left_indent = Inches(0.39375)
    p_item.paragraph_format.first_line_indent = Inches(-0.196527)
    run_item = p_item.add_run(item)
    format_text_run(run_item)

doc.save(dest_path)
print("Saved modified document successfully to", dest_path)
