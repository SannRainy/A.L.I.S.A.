# -*- coding: utf-8 -*-
"""
generate_actual_d_docx.py
=========================
Generates the revised Section D. Pengujian Word document containing ACTUAL testing data
from the TVJP automated test suite (Black-box, White-box, PBT, MT).
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_table_horizontal_borders(table, border_color="000000", inside_color="D3D3D3"):
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="12" w:space="0" w:color="{border_color}"/>
        <w:bottom w:val="single" w:sz="12" w:space="0" w:color="{border_color}"/>
        <w:left w:val="none"/>
        <w:right w:val="none"/>
        <w:insideH w:val="single" w:sz="6" w:space="0" w:color="{inside_color}"/>
        <w:insideV w:val="none"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))

def set_table_cell_margins(table, top=100, bottom=100, left=120, right=120):
    tblPr = table._tbl.tblPr
    margins_xml = f'''
    <w:tblCellMar {nsdecls("w")}>
        <w:top w:w="{top}" w:type="dxa"/>
        <w:bottom w:w="{bottom}" w:type="dxa"/>
        <w:left w:w="{left}" w:type="dxa"/>
        <w:right w:w="{right}" w:type="dxa"/>
    </w:tblCellMar>
    '''
    tblPr.append(parse_xml(margins_xml))

def format_word_cell(cell, text: str, bold=False, italic=False, color=None, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(str(text))
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color

def create_document():
    doc = docx.Document()
    
    # 1 inch margin
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title Style
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    def add_paragraph(text, italic=False, justify=True, space_after=12):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(space_after)
        if justify:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.italic = italic

    # Content
    add_heading_1("D. Pengujian Sistem")
    
    add_paragraph(
        "Bagian ini memaparkan hasil pengujian komprehensif pada sistem Tutor Virtual Bahasa Jepang (TVJP) "
        "A.L.I.S.A. Pengujian dilakukan dengan membagi pengujian ke dalam empat metodologi utama yang saling "
        "melengkapi: Black-box Testing (pada level sistem E2E dan UAT), White-box Testing (pada level unit "
        "dan logika percabangan), Property-Based Testing (PBT), dan Metamorphic Testing (MT) untuk mengatasi "
        "Oracle Problem pada komponen kecerdasan buatan (AI) di sistem."
    )

    # 1. Black-box testing
    add_heading_2("1. Black-box testing")
    add_paragraph(
        "Pengujian Black-box berfokus pada evaluasi fungsionalitas sistem dari perspektif eksternal (input-output) "
        "tanpa menganalisis alur kode internal. Pengujian ini diterapkan pada dua level, yaitu System Testing "
        "dan User Acceptance Testing (UAT)."
    )

    # 1.1 E2E System Testing (Layer 3)
    add_heading_3("1.1 E2E System Testing")
    add_paragraph(
        "Pengujian sistem berfokus pada alur fungsional ujung-ke-ujung (end-to-end) lintas modul, memastikan "
        "integrasi antara mesin BKT, SRS, sistem streak, dan pemrosesan audio berjalan dengan mulus."
    )
    
    add_paragraph("Tabel 4.1 Hasil Pengujian E2E System Testing (Black-Box)", italic=True, space_after=4)

    # Table 4.1 Data
    t1_headers = ["No", "Skenario", "Case", "Data", "Hasil yang Diharapkan", "Hasil Pengujian", "Kesimpulan"]
    t1_cols = [Inches(0.4), Inches(1.3), Inches(1.2), Inches(1.2), Inches(1.4), Inches(0.9), Inches(0.8)]
    t1_rows = [
        {
            "no": "1",
            "skenario": "S-01: E2E Learning Session",
            "case": "Sesi belajar kosakata normal dengan input jawaban benar",
            "data": "User submit jawaban benar",
            "expected": "P(L) > 0.500 | streak_days >= 1",
            "actual": "P(L)=0.823 | streak_days=1",
            "status": "[√] Valid"
        },
        {
            "no": "2",
            "skenario": "S-02: E2E Audio Conversation",
            "case": "Alur input suara user hingga tutor memberikan balasan suara",
            "data": "File WAV percakapan user",
            "expected": "Audio Output: '*.wav'",
            "actual": "Audio Output: 'response_1782.wav'",
            "status": "[√] Valid"
        },
        {
            "no": "3",
            "skenario": "S-03: Adaptive Question Selector",
            "case": "Pemilihan soal adaptif berdasarkan tingkat kemahiran kognitif",
            "data": "beliefs={n1:0.4, n2:0.9} -> benar",
            "expected": "Next_Node = 'n1' | P(L) > 0.400",
            "actual": "Next_Node = 'n1' | P(L)=0.759",
            "status": "[√] Valid"
        }
    ]
    
    # Render Table 4.1
    table1 = doc.add_table(rows=1 + len(t1_rows), cols=7)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_horizontal_borders(table1)
    set_table_cell_margins(table1)
    
    # Header
    for idx, name in enumerate(t1_headers):
        cell = table1.rows[0].cells[idx]
        cell.width = t1_cols[idx]
        set_cell_background(cell, "D3D3D3")
        format_word_cell(cell, name, bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        
    for r_idx, row in enumerate(t1_rows):
        cells = table1.rows[r_idx+1].cells
        for col_idx in range(7):
            cells[col_idx].width = t1_cols[col_idx]
        format_word_cell(cells[0], row["no"], font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        format_word_cell(cells[1], row["skenario"], font_size=9)
        format_word_cell(cells[2], row["case"], font_size=9)
        format_word_cell(cells[3], row["data"], font_size=9)
        format_word_cell(cells[4], row["expected"], font_size=9)
        format_word_cell(cells[5], row["actual"], font_size=9)
        format_word_cell(cells[6], row["status"], font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph().paragraph_format.space_before = Pt(6)

    # 1.2 User Acceptance Testing (Layer 4)
    add_heading_3("1.2 User Acceptance Testing (UAT)")
    add_paragraph(
        "UAT dilakukan untuk menguji apakah sistem dapat memenuhi skenario kebutuhan belajar riil siswa, "
        "seperti motivasi streak, pembagian tingkat kesulitan kuis, sistem penghargaan XP, dan umpan balik kesalahan."
    )
    
    add_paragraph("Tabel 4.2 Hasil Pengujian User Acceptance Testing (Black-Box)", italic=True, space_after=4)

    # Table 4.2 Data
    t2_rows = [
        {
            "no": "1",
            "skenario": "A-01: Appropriate Difficulty Allocation",
            "case": "Pemberian tingkat kesulitan kuis sesuai kompetensi siswa",
            "data": "Profil siswa P(L) = 0.20",
            "expected": "Label kesulitan = 'easy'",
            "actual": "Label kesulitan = 'easy'",
            "status": "[√] Valid"
        },
        {
            "no": "2",
            "skenario": "A-02: Streak Tracker Motivation",
            "case": "Pencatatan konsistensi belajar harian pengguna secara runtut",
            "data": "Aktivitas hari ini dan kemarin",
            "expected": "streak_days = 2",
            "actual": "streak_days = 2",
            "status": "[√] Valid"
        },
        {
            "no": "3",
            "skenario": "A-03: Gamification XP Scaling",
            "case": "Akumulasi poin XP dan kenaikan tingkat level otomatis",
            "data": "XP awal = 95, aksi master (+10 XP)",
            "expected": "XP = 105 | Level = 2",
            "actual": "XP = 105 | Level = 2",
            "status": "[√] Valid"
        },
        {
            "no": "4",
            "skenario": "A-04: Actionable Grammar Feedback",
            "case": "Penayangan pesan koreksi tata bahasa yang salah tik",
            "data": "Kalimat: '私はりんgoをを食べる'",
            "expected": "Feedback terstruktur (助詞重複)",
            "actual": "Error feedback: '助詞「を」が重複...'",
            "status": "[√] Valid"
        }
    ]

    # Render Table 4.2
    table2 = doc.add_table(rows=1 + len(t2_rows), cols=7)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_horizontal_borders(table2)
    set_table_cell_margins(table2)
    
    # Header
    for idx, name in enumerate(t1_headers):
        cell = table2.rows[0].cells[idx]
        cell.width = t1_cols[idx]
        set_cell_background(cell, "D3D3D3")
        format_word_cell(cell, name, bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        
    for r_idx, row in enumerate(t2_rows):
        cells = table2.rows[r_idx+1].cells
        for col_idx in range(7):
            cells[col_idx].width = t1_cols[col_idx]
        format_word_cell(cells[0], row["no"], font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        format_word_cell(cells[1], row["skenario"], font_size=9)
        format_word_cell(cells[2], row["case"], font_size=9)
        format_word_cell(cells[3], row["data"], font_size=9)
        format_word_cell(cells[4], row["expected"], font_size=9)
        format_word_cell(cells[5], row["actual"], font_size=9)
        format_word_cell(cells[6], row["status"], font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph().paragraph_format.space_before = Pt(6)

    # 2. White-box testing
    add_heading_2("2. White-box testing")
    add_paragraph(
        "Pengujian White-box berfokus pada kebenaran struktur program secara internal, menguji logika modular "
        "(unit testing) serta validasi jalur kondisi batas (logic boundary testing)."
    )

    # 2.1 Unit Testing (Layer 1)
    add_heading_3("2.1 Unit Testing")
    add_paragraph(
        "Unit testing memvalidasi logika paling dasar dari modul-modul sistem, seperti perhitungan matematika "
        "BKT, LLM intents detection, penyaringan teks untuk TTS, tokenisasi morfologi, serta rumus SM-2."
    )
    
    add_paragraph("Tabel 4.3 Hasil Pengujian Unit Testing (White-Box)", italic=True, space_after=4)

    # Table 4.3 Data
    t3_rows = [
        {
            "no": "1",
            "skenario": "U-01: BKT CAT Difficulty",
            "case": "Penentuan kesulitan adaptif CAT berdasarkan belief",
            "data": "p_mastered = 0.85, 0.50, 0.20",
            "expected": "0.85->hard | 0.50->medium | 0.20->easy",
            "actual": "0.85->hard | 0.50->medium | 0.20->easy",
            "status": "[√] Valid"
        },
        {
            "no": "2",
            "skenario": "U-02: BKT Bayesian Update",
            "case": "Pembaruan probabilistik BKT saat benar vs salah",
            "data": "p_l=0.5, correct=True / False",
            "expected": "Benar->P(L)~0.8229 | Salah->P(L)~0.2031",
            "actual": "Benar->P(L)=0.8229 | Salah->P(L)=0.2031",
            "status": "[√] Valid"
        },
        {
            "no": "3",
            "skenario": "U-06: LLM Quiz Intent",
            "case": "Deteksi intensitas instruksi kuis pada masukan prompt",
            "data": "text = 'Tolong berikan latihan kuis N5'",
            "expected": "is_quiz = True",
            "actual": "is_quiz = True",
            "status": "[√] Valid"
        },
        {
            "no": "4",
            "skenario": "U-12: Double Particle Detect",
            "case": "Pendeteksian duplikasi partikel gramatikal bahasa Jepang",
            "data": "text = '私はりんごをを食べる'",
            "expected": "Terdeteksi particle_error (を)",
            "actual": "Terdeteksi particle_error (を)",
            "status": "[√] Valid"
        },
        {
            "no": "5",
            "skenario": "U-13: SRS SM-2 Pure Math",
            "case": "Kalkulasi interval dan EF awal repetisi SM-2",
            "data": "quality=5, rep=0",
            "expected": "interval=1, EF=2.60",
            "actual": "interval=1, EF=2.60",
            "status": "[√] Valid"
        }
    ]

    # Render Table 4.3
    table3 = doc.add_table(rows=1 + len(t3_rows), cols=7)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_horizontal_borders(table3)
    set_table_cell_margins(table3)
    
    # Header
    for idx, name in enumerate(t1_headers):
        cell = table3.rows[0].cells[idx]
        cell.width = t1_cols[idx]
        set_cell_background(cell, "D3D3D3")
        format_word_cell(cell, name, bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        
    for r_idx, row in enumerate(t3_rows):
        cells = table3.rows[r_idx+1].cells
        for col_idx in range(7):
            cells[col_idx].width = t1_cols[col_idx]
        format_word_cell(cells[0], row["no"], font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        format_word_cell(cells[1], row["skenario"], font_size=9)
        format_word_cell(cells[2], row["case"], font_size=9)
        format_word_cell(cells[3], row["data"], font_size=9)
        format_word_cell(cells[4], row["expected"], font_size=9)
        format_word_cell(cells[5], row["actual"], font_size=9)
        format_word_cell(cells[6], row["status"], font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph().paragraph_format.space_before = Pt(6)

    # 2.2 Logic & Boundary Testing (Layer 5)
    add_heading_3("2.2 Logic & Boundary Testing")
    add_paragraph(
        "Pengujian logika struktural bertujuan memverifikasi apakah cabang kondisi kaku dan nilai ekstrem batas "
        "pada algoritma BKT, pembatasan EF floor SM-2, dan toleransi gap streak berjalan dengan benar."
    )
    
    add_paragraph("Tabel 4.4 Hasil Pengujian Logic & Boundary Testing (White-Box)", italic=True, space_after=4)

    # Table 4.4 Data
    t4_rows = [
        {
            "no": "1",
            "skenario": "L-01: BKT Boundary Precision",
            "case": "Validasi cabang logika transisi kategori kemampun kaku",
            "data": "p_l = 0.849 vs p_l = 0.850",
            "expected": "0.849->Medium (n_below) | 0.850->Hard (n_high)",
            "actual": "0.849->Medium | 0.850->Hard",
            "status": "[√] Valid"
        },
        {
            "no": "2",
            "skenario": "L-02: SM-2 Algorithmic Constraint",
            "case": "Penguncian nilai minimum EF agar tidak di bawah floor",
            "data": "Salah 10x berturut-turut (Quality=0)",
            "expected": "EF tertahan di floor = 1.30",
            "actual": "Easiness Factor Floor = 1.30",
            "status": "[√] Valid"
        },
        {
            "no": "3",
            "skenario": "L-03: Full Probability Branch",
            "case": "Cakupan boundary probability clamping BKT",
            "data": "p_prior = 0.999, correct = True",
            "expected": "Output ter-clamp aman di max 0.999",
            "actual": "Matrix outcome = 0.9990",
            "status": "[√] Valid"
        },
        {
            "no": "4",
            "skenario": "L-04: Streak Gap Absense Forgive",
            "case": "Kondisi batas toleransi bolos 1 hari vs reset streak",
            "data": "Bolos 1 hari vs Bolos 2 hari berturut",
            "expected": "Bolos 1 hari -> Gap1 (Toleransi) | 2 hari -> Reset",
            "actual": "Gap1 = 1 (Toleransi) | Gap2 = 0 (Reset)",
            "status": "[√] Valid"
        }
    ]

    # Render Table 4.4
    table4 = doc.add_table(rows=1 + len(t4_rows), cols=7)
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_horizontal_borders(table4)
    set_table_cell_margins(table4)
    
    # Header
    for idx, name in enumerate(t1_headers):
        cell = table4.rows[0].cells[idx]
        cell.width = t1_cols[idx]
        set_cell_background(cell, "D3D3D3")
        format_word_cell(cell, name, bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        
    for r_idx, row in enumerate(t4_rows):
        cells = table4.rows[r_idx+1].cells
        for col_idx in range(7):
            cells[col_idx].width = t1_cols[col_idx]
        format_word_cell(cells[0], row["no"], font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        format_word_cell(cells[1], row["skenario"], font_size=9)
        format_word_cell(cells[2], row["case"], font_size=9)
        format_word_cell(cells[3], row["data"], font_size=9)
        format_word_cell(cells[4], row["expected"], font_size=9)
        format_word_cell(cells[5], row["actual"], font_size=9)
        format_word_cell(cells[6], row["status"], font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph().paragraph_format.space_before = Pt(6)

    # 3. Property-Based Testing (PBT)
    add_heading_2("3. Property-Based Testing (PBT)")
    add_paragraph(
        "Property-Based Testing (PBT) menggunakan pustaka Hypothesis untuk memverifikasi invarian "
        "(sifat/hukum sistem yang harus selalu benar) dengan melakukan pengujian fungsionalitas fuzzer otomatis "
        "menggunakan ratusan data acak masukan."
    )
    
    add_paragraph("Tabel 4.5 Hasil Pengujian Property-Based Testing (PBT)", italic=True, space_after=4)

    # Table 4.5 Data
    t5_rows = [
        {
            "no": "1",
            "skenario": "PBT-01: Crash Safety Hiragana",
            "case": "Invarian crash safety NLP konversi Romaji Hiragana",
            "data": "Fuzzing acak Hiragana (~300 sampel)",
            "expected": "Hasil selalu bertipe string & tidak crash",
            "actual": "Hasil valid tipe str (Zero exceptions)",
            "status": "[√] Valid"
        },
        {
            "no": "2",
            "skenario": "PBT-03: Empty Input Invariant",
            "case": "Pencegahan error NLP pada input string kosong/blank",
            "data": "Fuzzing acak spasi & newline (~100 sampel)",
            "expected": "Return string kosong tanpa exception",
            "actual": "Invarian kosong terjaga (No exception)",
            "status": "[√] Valid"
        },
        {
            "no": "3",
            "skenario": "PBT-05: SM-2 EF Floor Invariant",
            "case": "Batas bawah invariant EF harus >= 1.30",
            "data": "Fuzzing initial EF & rating Q (~500 sampel)",
            "expected": "Output EF selalu >= 1.30",
            "actual": "EF >= 1.30 terjaga 100%",
            "status": "[√] Valid"
        },
        {
            "no": "4",
            "skenario": "PBT-06: SM-2 Interval Invariant",
            "case": "Invarian nilai interval review harus berupa bilangan positif",
            "data": "Fuzzing kombinasi Q & EF (~400 sampel)",
            "expected": "interval_days selalu >= 1 hari",
            "actual": "interval_days >= 1 terjaga 100%",
            "status": "[√] Valid"
        },
        {
            "no": "5",
            "skenario": "PBT-07: BKT Boundary Invariant",
            "case": "Invarian nilai probabilitas BKT dalam batas valid",
            "data": "Fuzzing prior probability acak (~500 sampel)",
            "expected": "Output belief selalu dalam range [0.001, 0.999]",
            "actual": "Output belief bounded [0.001, 0.999]",
            "status": "[√] Valid"
        }
    ]

    # Render Table 4.5
    table5 = doc.add_table(rows=1 + len(t5_rows), cols=7)
    table5.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_horizontal_borders(table5)
    set_table_cell_margins(table5)
    
    # Header
    for idx, name in enumerate(t1_headers):
        cell = table5.rows[0].cells[idx]
        cell.width = t1_cols[idx]
        set_cell_background(cell, "D3D3D3")
        format_word_cell(cell, name, bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        
    for r_idx, row in enumerate(t5_rows):
        cells = table5.rows[r_idx+1].cells
        for col_idx in range(7):
            cells[col_idx].width = t1_cols[col_idx]
        format_word_cell(cells[0], row["no"], font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        format_word_cell(cells[1], row["skenario"], font_size=9)
        format_word_cell(cells[2], row["case"], font_size=9)
        format_word_cell(cells[3], row["data"], font_size=9)
        format_word_cell(cells[4], row["expected"], font_size=9)
        format_word_cell(cells[5], row["actual"], font_size=9)
        format_word_cell(cells[6], row["status"], font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph().paragraph_format.space_before = Pt(6)

    # 4. Metamorphic Testing (MT)
    add_heading_2("4. Metamorphic Testing (MT)")
    add_paragraph(
        "Metamorphic Testing (MT) diimplementasikan untuk memecahkan Oracle Problem pada komponen sistem yang berbasis "
        "non-deterministik/AI. Metode ini mengevaluasi kebenaran keluaran dengan memverifikasi hubungan relasi "
        "metamorfik (Metamorphic Relations atau MR) di antara beberapa eksekusi masukan yang saling terkait."
    )
    
    add_paragraph("Tabel 4.6 Hasil Pengujian Metamorphic Testing (MT)", italic=True, space_after=4)

    # Table 4.6 Data
    t6_rows = [
        {
            "no": "1",
            "skenario": "MT-01: MR-1 BKT Monotonic Increase",
            "case": "Relasi: Prioritas belief harus meningkat jika jawaban benar",
            "data": "f(p, True) vs p",
            "expected": "f(p, True) > p",
            "actual": "Terbukti untuk 10 cases (e.g. 0.7120 > 0.4000)",
            "status": "[√] Valid"
        },
        {
            "no": "2",
            "skenario": "MT-02: MR-1b BKT Comparative Decrease",
            "case": "Relasi: Output belief jawaban salah harus lebih kecil dari jawaban benar",
            "data": "f(p, False) vs f(p, True)",
            "expected": "f(p, False) < f(p, True)",
            "actual": "Terbukti 11 titik (e.g. 0.2031 < 0.8229)",
            "status": "[√] Valid"
        },
        {
            "no": "3",
            "skenario": "MT-04: MR-2 SM-2 Interval Monotonic",
            "case": "Relasi: Kualitas jawaban lebih tinggi harus menghasilkan interval lebih panjang",
            "data": "q3 vs q4 pada rep=2",
            "expected": "Interval(q4) >= Interval(q3)",
            "actual": "Terbukti 288 kasus grid (15 >= 12)",
            "status": "[√] Valid"
        },
        {
            "no": "4",
            "skenario": "MT-05: MR-3 SM-2 Reset Invariant",
            "case": "Relasi: Quality kegagalan (Q < 3) harus mereset repetisi ke 0",
            "data": "quality = 0, 1, 2 pada rep=10",
            "expected": "rep == 0 dan interval == 1",
            "actual": "Reset invariant terjaga 45 combo (rep=0, int=1)",
            "status": "[√] Valid"
        },
        {
            "no": "5",
            "skenario": "MT-07: MR-6 SM-2 EF Monotonicity",
            "case": "Relasi: Kualitas jawaban lebih tinggi menghasilkan kenaikan EF lebih besar",
            "data": "q5 vs q4",
            "expected": "EF(q5) >= EF(q4)",
            "actual": "Terbukti 180 pasangan (2.70 >= 2.60)",
            "status": "[√] Valid"
        }
    ]

    # Render Table 4.6
    table6 = doc.add_table(rows=1 + len(t6_rows), cols=7)
    table6.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_horizontal_borders(table6)
    set_table_cell_margins(table6)
    
    # Header
    for idx, name in enumerate(t1_headers):
        cell = table6.rows[0].cells[idx]
        cell.width = t1_cols[idx]
        set_cell_background(cell, "D3D3D3")
        format_word_cell(cell, name, bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        
    for r_idx, row in enumerate(t6_rows):
        cells = table6.rows[r_idx+1].cells
        for col_idx in range(7):
            cells[col_idx].width = t1_cols[col_idx]
        format_word_cell(cells[0], row["no"], font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        format_word_cell(cells[1], row["skenario"], font_size=9)
        format_word_cell(cells[2], row["case"], font_size=9)
        format_word_cell(cells[3], row["data"], font_size=9)
        format_word_cell(cells[4], row["expected"], font_size=9)
        format_word_cell(cells[5], row["actual"], font_size=9)
        format_word_cell(cells[6], row["status"], font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Save Document
    output_fn = "D_Pengujian_TVJP.docx"
    doc.save(output_fn)
    print(f"Success: Generated {output_fn}")

if __name__ == "__main__":
    create_document()
