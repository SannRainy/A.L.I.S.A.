# -*- coding: utf-8 -*-
"""
generate_final_docx.py
======================
Generates the final Section D. Pengujian Word document containing:
1. Black-box Testing: 11 pages (Login, Register, Dashboard, CSV, Chatbot, Quiz, Kanji, SRS, Speaking, Profile, Achievement)
   Using test accounts: Krisnasatyaarisandy@gmail.com (Admin), Raincallin@gmail.com (user), testing@gmail.com (user), Satyakrisna89=0@gmail.com (user)
2. White-box Testing (Unit & Logic)
3. Property-Based Testing (PBT)
4. Metamorphic Testing (MT)
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_table_horizontal_borders(table, border_color="000000", inside_color="D3D3D3"):
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="12" w:space="0" w:color="{border_color}"/>
        <w:bottom w:val="single" w:sz="12" w:space="0" w:color="{border_color}"/>
        <w:left w:val="none"/>
        <w:right w:val="none"/>
        <w:insideH w:val="single" w:sz="6" w:space="0" w:color="{inside_color}"/>
        <w:insideV w:val="none"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))

def set_table_cell_margins(table, top=100, bottom=100, left=120, right=120):
    tblPr = table._tbl.tblPr
    margins_xml = f'''
    <w:tblCellMar {nsdecls("w")}>
        <w:top w:w="{top}" w:type="dxa"/>
        <w:bottom w:w="{bottom}" w:type="dxa"/>
        <w:left w:w="{left}" w:type="dxa"/>
        <w:right w:w="{right}" w:type="dxa"/>
    </w:tblCellMar>
    '''
    tblPr.append(parse_xml(margins_xml))

def format_word_cell(cell, text: str, bold=False, italic=False, color=None, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(str(text))
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color

def create_document():
    doc = docx.Document()
    
    # Page setup (1 inch margin)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styles
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    def add_paragraph(text, italic=False, justify=True, space_after=12):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(space_after)
        if justify:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.italic = italic

    # Common Table Renderer
    headers = ["No", "Skenario", "Case", "Data", "Hasil yang Diharapkan", "Hasil Pengujian", "Kesimpulan"]
    col_widths = [Inches(0.4), Inches(1.2), Inches(1.3), Inches(1.2), Inches(1.4), Inches(0.9), Inches(0.8)]

    def render_table(caption, rows_data):
        add_paragraph(caption, italic=True, space_after=4)
        table = doc.add_table(rows=1 + len(rows_data), cols=7)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_horizontal_borders(table)
        set_table_cell_margins(table)

        # Header
        for i, text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.width = col_widths[i]
            set_cell_background(cell, "D3D3D3")
            format_word_cell(cell, text, bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Body Rows
        for r_idx, row in enumerate(rows_data):
            row_cells = table.rows[r_idx + 1].cells
            for c_idx in range(7):
                row_cells[c_idx].width = col_widths[c_idx]
            format_word_cell(row_cells[0], row["no"], font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            format_word_cell(row_cells[1], row["skenario"], font_size=9)
            format_word_cell(row_cells[2], row["case"], font_size=9)
            format_word_cell(row_cells[3], row["data"], font_size=9)
            format_word_cell(row_cells[4], row["expected"], font_size=9)
            format_word_cell(row_cells[5], row["actual"], font_size=9)
            format_word_cell(row_cells[6], row["status"], font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Add spacing paragraph after table
        doc.add_paragraph().paragraph_format.space_before = Pt(6)

    def render_table_5(caption, rows_data):
        add_paragraph(caption, italic=True, space_after=4)
        headers_5 = ["Path", "Skenario Uji", "Kondisi Input", "Expected Output", "Status"]
        col_widths_5 = [Inches(0.6), Inches(1.8), Inches(1.8), Inches(1.5), Inches(0.8)]
        table = doc.add_table(rows=1 + len(rows_data), cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_horizontal_borders(table)
        set_table_cell_margins(table)

        # Header
        for i, text in enumerate(headers_5):
            cell = table.rows[0].cells[i]
            cell.width = col_widths_5[i]
            set_cell_background(cell, "D3D3D3")
            format_word_cell(cell, text, bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Body Rows
        for r_idx, row in enumerate(rows_data):
            row_cells = table.rows[r_idx + 1].cells
            for c_idx in range(5):
                row_cells[c_idx].width = col_widths_5[c_idx]
            format_word_cell(row_cells[0], row["path"], font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            format_word_cell(row_cells[1], row["skenario"], font_size=9)
            format_word_cell(row_cells[2], row["input"], font_size=9)
            format_word_cell(row_cells[3], row["expected"], font_size=9)
            format_word_cell(row_cells[4], row["status"], font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Add spacing paragraph after table
        doc.add_paragraph().paragraph_format.space_before = Pt(6)

    # DOCUMENT TEXT GENERATION
    add_heading_1("D. Pengujian Sistem")
    
    add_paragraph(
        "Bagian ini memaparkan hasil pengujian komprehensif pada sistem Tutor Virtual Bahasa Jepang (TVJP) "
        "A.L.I.S.A. Pengujian dilakukan dengan membagi pengujian ke dalam empat metodologi utama yang saling "
        "melengkapi: Black-box Testing (pada level antarmuka halaman pengguna), White-box Testing (pada level unit "
        "dan logika percabangan backend), Property-Based Testing (PBT), dan Metamorphic Testing (MT) untuk mengatasi "
        "Oracle Problem pada komponen kecerdasan buatan (AI) di sistem."
    )

    # 1. Black-box testing
    add_heading_2("1. Black-box testing")
    add_paragraph(
        "Pengujian Black-box berfokus pada evaluasi fungsionalitas antarmuka aplikasi dari sudut pandang "
        "pengguna (input-output) tanpa menganalisis alur kode internal. Pengujian ini dilakukan per halaman "
        "pada 11 antarmuka utama sistem dengan menggunakan data akun pengujian riil: Krisnasatyaarisandy@gmail.com (Admin), "
        "Raincallin@gmail.com (user), testing@gmail.com (user), dan Satyakrisna89=0@gmail.com (user)."
    )

    # 1.1 Halaman Login
    add_heading_3("1.1 Testing Halaman Login")
    add_paragraph("Pengujian halaman login dilakukan untuk memverifikasi fungsionalitas autentikasi akun pengguna dan administrator masuk ke dalam sistem.")
    login_data = [
        {
            "no": "1", "skenario": "Login sukses user umum",
            "case": "1. Input email user\n2. Input password valid\n3. Klik tombol Masuk Sekarang",
            "data": "Email: Raincallin@gmail.com\nPassword: Password123 (valid)",
            "expected": "Sistem menampilkan toast sukses 'Berhasil masuk! Selamat datang kembali.' dan mengarahkan ke homepage chatbot.",
            "actual": "Berhasil sesuai harapan, toast muncul dan homepage termuat.", "status": "[√] Valid"
        },
        {
            "no": "2", "skenario": "Login sukses admin",
            "case": "1. Input email admin\n2. Input password admin\n3. Klik tombol Masuk Sekarang",
            "data": "Email: Krisnasatyaarisandy@gmail.com\nPassword: PasswordAdmin123",
            "expected": "Sistem menampilkan toast 'Selamat datang, Admin! Mengalihkan...' dan mengarahkan ke /admin.",
            "actual": "Berhasil sesuai harapan, rute admin /admin terbuka penuh.", "status": "[√] Valid"
        },
        {
            "no": "3", "skenario": "Login gagal karena password salah",
            "case": "1. Input email user\n2. Input password salah\n3. Klik tombol Masuk Sekarang",
            "data": "Email: testing@gmail.com\nPassword: PasswordSalah",
            "expected": "Sistem menampilkan pesan error 'Login gagal. Periksa email dan password' di bawah form login.",
            "actual": "Berhasil sesuai harapan, error merah muncul di form login.", "status": "[√] Valid"
        },
        {
            "no": "4", "skenario": "Login gagal karena format email tidak valid",
            "case": "1. Input email tanpa '@'\n2. Klik tombol Masuk Sekarang",
            "data": "Email: testingemail.com\nPassword: Password123",
            "expected": "Sistem menampilkan pesan error lokal 'Format email tidak valid.' di bawah form.",
            "actual": "Berhasil sesuai harapan, validasi lokal memblokir request.", "status": "[√] Valid"
        }
    ]
    render_table("Tabel 4.1 Testing Halaman Login", login_data)

    # 1.2 Halaman Register
    add_heading_3("1.2 Testing Halaman Register")
    add_paragraph("Pengujian halaman register dilakukan untuk memverifikasi pendaftaran akun baru pengguna beserta penyimpanan data demografis belajar.")
    register_data = [
        {
            "no": "1", "skenario": "Registrasi sukses pengguna baru",
            "case": "1. Input formulir akun baru\n2. Lengkapi data demografis\n3. Klik tombol Daftar Sekarang",
            "data": "Email: Satyakrisna89=0@gmail.com\nPass: Password123\nNama: Satya Krisna\nUmur: 22, Gender: Laki-laki\nNegara: Indonesia\nTujuan: Hobi",
            "expected": "Akun terdaftar via POST /api/v1/auth/register, menampilkan toast 'Registrasi berhasil!', dan form isRegistering beralih ke false.",
            "actual": "Berhasil sesuai harapan, user baru terdaftar di database profiles.", "status": "[√] Valid"
        },
        {
            "no": "2", "skenario": "Registrasi gagal karena email sudah terdaftar",
            "case": "1. Input email terdaftar\n2. Lengkapi formulir\n3. Klik tombol Daftar Sekarang",
            "data": "Email: Raincallin@gmail.com\nPassword: Password123",
            "expected": "Sistem membatalkan pendaftaran dan menampilkan pesan error 'Registrasi gagal.' di bawah form.",
            "actual": "Berhasil sesuai harapan, error server tereduksi ke form.", "status": "[√] Valid"
        },
        {
            "no": "3", "skenario": "Registrasi gagal karena password kurang dari 6 karakter",
            "case": "1. Input password < 6 karakter\n2. Klik tombol Daftar Sekarang",
            "data": "Email: userbaru@gmail.com\nPassword: 12345",
            "expected": "Sistem memvalidasi secara lokal dan menampilkan error 'Password minimal 6 karakter.'.",
            "actual": "Berhasil sesuai harapan, validasi memblokir pendaftaran.", "status": "[√] Valid"
        },
        {
            "no": "4", "skenario": "Registrasi gagal karena negara tidak terdaftar",
            "case": "1. Ketik negara asal non-list\n2. Klik tombol Daftar Sekarang",
            "data": "Negara: Wakanda",
            "expected": "Sistem memvalidasi secara lokal dan menampilkan error 'Harap pilih asal negara yang valid dari daftar.'.",
            "actual": "Berhasil sesuai harapan, input negara diblokir.", "status": "[√] Valid"
        }
    ]
    render_table("Tabel 4.2 Testing Halaman Register", register_data)

    # 1.3 Halaman Dashboard Admin
    add_heading_3("1.3 Testing Halaman Dashboard Admin")
    add_paragraph("Pengujian halaman dashboard admin dilakukan untuk memastikan keamanan hak akses berbasis peran (RBAC) serta visualisasi grafik statistik belajar.")
    admin_data = [
        {
            "no": "1", "skenario": "Akses grafik analisis sebagai admin",
            "case": "1. Login dengan akun admin\n2. Masuk rute /admin\n3. Buka tab Analisis & Grafik",
            "data": "Email: Krisnasatyaarisandy@gmail.com\nRole: admin",
            "expected": "Halaman admin terbuka penuh dan merender sebaran demografis, skor kuis, dan grafik status AI models.",
            "actual": "Berhasil sesuai harapan, grafik termuat di AnalyticsTab.", "status": "[√] Valid"
        },
        {
            "no": "2", "skenario": "Membuka daftar pengguna di dashboard",
            "case": "1. Klik tab Users\n2. Amati daftar pengguna",
            "data": "Email: Krisnasatyaarisandy@gmail.com",
            "expected": "Sistem memuat dan merender tabel daftar semua pengguna beserta detail statistik belajar (Level, XP, Neo4j stats).",
            "actual": "Berhasil sesuai harapan, data termuat di UsersTab.", "status": "[√] Valid"
        },
        {
            "no": "3", "skenario": "Ekspor data pengguna ke CSV/JSON",
            "case": "1. Klik tombol Export CSV atau Export JSON",
            "data": "Format: CSV / JSON",
            "expected": "Browser mengunduh file 'tvjp_users_export.csv' atau 'tvjp_users_export.json' secara real-time.",
            "actual": "Berhasil sesuai harapan, file terunduh dengan benar.", "status": "[√] Valid"
        },
        {
            "no": "4", "skenario": "Akses halaman admin ditolak untuk user biasa",
            "case": "1. Login dengan akun user\n2. Akses rute /admin secara manual",
            "data": "Email: testing@gmail.com\nRole: user",
            "expected": "Sistem menolak akses dan mengalihkan pengguna kembali ke halaman utama (/).",
            "actual": "Berhasil sesuai harapan, rute dialihkan otomatis.", "status": "[√] Valid"
        }
    ]
    render_table("Tabel 4.3 Testing Halaman Dashboard Admin", admin_data)

    # 1.4 Halaman CSV File Management
    add_heading_3("1.4 Testing Halaman CSV File Management")
    add_paragraph("Pengujian halaman CSV File Management dilakukan untuk memverifikasi kelancaran pembaruan data kurikulum dan sinkronisasi ke Neo4j.")
    csv_data = [
        {
            "no": "1", "skenario": "Mengubah data kosakata",
            "case": "1. Masuk ke tab Data Pipeline\n2. Klik berkas vocab.csv\n3. Ubah arti kosakata di tabel grid\n4. Klik Save CSV",
            "data": "Email: Krisnasatyaarisandy@gmail.com\nMengubah kosakata '食べる' -> 'Makan'",
            "expected": "Data berhasil diperbarui di server via API PUT, menampilkan alert 'File CSV berhasil disimpan'.",
            "actual": "Berhasil sesuai harapan, alert muncul dan perubahan tersimpan.", "status": "[√] Valid"
        },
        {
            "no": "2", "skenario": "Sinkronisasi otomatis ke database graf Neo4j",
            "case": "1. Klik tab Data Pipeline\n2. Klik tombol Trigger Ingest",
            "data": "Memicu API POST /ingest",
            "expected": "Server backend FastAPI memproses berkas CSV menggunakan ingest_n5.py dan memutakhirkan simpul dan relasi Neo4j.",
            "actual": "Berhasil sesuai harapan, log output Neo4j ingestion muncul.", "status": "[√] Valid"
        }
    ]
    render_table("Tabel 4.4 Testing Halaman CSV File Management", csv_data)

    # 1.5 Homepage Chatbot Mode
    add_heading_3("1.5 Testing Homepage Chatbot Mode")
    add_paragraph("Pengujian chatbot mode dilakukan untuk menguji fitur obrolan teks bebas, teks furigana Jepang, serta kegunaan chip saran topik.")
    chatbot_data = [
        {
            "no": "1", "skenario": "Mengajukan pertanyaan tata bahasa",
            "case": "1. Tulis pertanyaan di input chat\n2. Klik tombol kirim",
            "data": "Email: Raincallin@gmail.com\nPertanyaan: 'Tolong jelaskan partikel は'",
            "expected": "Pesan dikirim via WS, avatar 3D Alisa menunjukkan animasi berbicara, balasan mengalir (streaming) dengan furigana (<ruby>) pada Kanji.",
            "actual": "Berhasil sesuai harapan, furigana dan animasi avatar aktif.", "status": "[√] Valid"
        },
        {
            "no": "2", "skenario": "Menggunakan Suggestion Chips",
            "case": "1. Amati chip saran di bawah obrolan\n2. Klik chip saran terpilih",
            "data": "Memicu klik chip: 'Bisa berikan contoh kalimat?'",
            "expected": "Pesan dari chip terkirim otomatis dan mendapat balasan berupa contoh kalimat dari tutor virtual.",
            "actual": "Berhasil sesuai harapan, chip terkirim dan dibalas Alisa.", "status": "[√] Valid"
        },
        {
            "no": "3", "skenario": "Menghapus riwayat percakapan (Clear Chat)",
            "case": "1. Klik tombol bersihkan riwayat (🗑️)",
            "data": "Klik tombol clear",
            "expected": "Riwayat chat terhapus dari store lokal dan area chat dibersihkan.",
            "actual": "Berhasil sesuai harapan, area chat kembali bersih.", "status": "[√] Valid"
        }
    ]
    render_table("Tabel 4.5 Testing Homepage Chatbot Mode", chatbot_data)

    # 1.6 Homepage Quiz Mode
    add_heading_3("1.6 Testing Homepage Quiz Mode")
    add_paragraph("Pengujian quiz mode dilakukan untuk memverifikasi peta tingkat level kuis sekuensial, interaksi pengerjaan latihan, dan perekaman data BKT.")
    quiz_data = [
        {
            "no": "1", "skenario": "Membuka peta materi sekuensial",
            "case": "1. Klik tab Quest Mode\n2. Amati status simpul peta belajar",
            "data": "Email: testing@gmail.com\nUser memiliki progres level 1",
            "expected": "Visualisasi peta kuis dirender, level 1 terbuka, sedangkan level berikutnya terkunci dengan ikon gembok.",
            "actual": "Berhasil sesuai harapan, peta quest termuat di QuestMap.", "status": "[√] Valid"
        },
        {
            "no": "2", "skenario": "Mengerjakan kuis QuestEngine",
            "case": "1. Klik level kuis terbuka\n2. Jawab semua soal kuis pilihan ganda hingga selesai",
            "data": "Menyelesaikan kuis level 1",
            "expected": "Menampilkan kuis. Setelah selesai, skor dan XP terhitung, serta data terkirim via POST /api/v1/quest/submit.",
            "actual": "Berhasil sesuai harapan, BKT kognitif user diperbarui.", "status": "[√] Valid"
        },
        {
            "no": "3", "skenario": "Pembatasan akses level terkunci",
            "case": "1. Klik level kuis yang terkunci (Level 2 ketika Level 1 belum tuntas)",
            "data": "Klik level terkunci",
            "expected": "Sistem menampilkan pop-up alert bertuliskan peringatan prasyarat tidak terpenuhi.",
            "actual": "Berhasil sesuai harapan, pop-up alert muncul.", "status": "[√] Valid"
        }
    ]
    render_table("Tabel 4.6 Testing Homepage Quiz Mode", quiz_data)

    # 1.7 Halaman Kanji Dojo
    add_heading_3("1.7 Testing Halaman Kanji Dojo")
    add_paragraph("Pengujian Kanji Dojo dilakukan untuk memverifikasi pembalikan flashcard kanji secara 3D, pemutaran suara lafal, dan animasi urutan goresan.")
    kanji_data = [
        {
            "no": "1", "skenario": "Mempelajari arti dan bacaan kanji N5",
            "case": "1. Pilih Kanji Dojo\n2. Pilih Set Kanji\n3. Klik flashcard kanji",
            "data": "Email: Raincallin@gmail.com\nFlashcard Kanji: '水'",
            "expected": "Kartu berputar dengan animasi 3D menampilkan cara baca onyomi/kunyomi dan terjemahan.",
            "actual": "Berhasil sesuai harapan, flip card berputar halus.", "status": "[√] Valid"
        },
        {
            "no": "2", "skenario": "Mengerjakan kuis evaluasi set Kanji",
            "case": "1. Selesaikan study flashcard\n2. Ketik jawaban arti pada kuis kanji\n3. Klik Cek",
            "data": "Input jawaban: 'Air'",
            "expected": "Sistem mencocokkan jawaban. Jika benar, memberikan feedback hijau dan menambah skor kuis.",
            "actual": "Berhasil sesuai harapan, feedback benar/salah muncul presisi.", "status": "[√] Valid"
        }
    ]
    render_table("Tabel 4.7 Testing Halaman Kanji Dojo", kanji_data)

    # 1.8 Halaman SRS Review
    add_heading_3("1.8 Testing Halaman SRS Review")
    add_paragraph("Pengujian SRS Review dilakukan untuk memvalidasi pemuatan antrean kartu hafalan terjadwal dan perhitungan ulasan ulang berdasar rating SM-2.")
    srs_data = [
        {
            "no": "1", "skenario": "Membuka antrean ulasan memori",
            "case": "1. Klik SRS Review Dojo dari QuestMap\n2. Amati kartu ulasan",
            "data": "Email: Satyakrisna89=0@gmail.com",
            "expected": "Memanggil API /api/v1/srs/due/{user_id}, menampilkan kartu review dan indikator progress.",
            "actual": "Berhasil sesuai harapan, daftar due items termuat.", "status": "[√] Valid"
        },
        {
            "no": "2", "skenario": "Menilai tingkat hafalan materi",
            "case": "1. Klik Tampilkan Jawaban\n2. Klik tombol rating memori (0-5)",
            "data": "Memilih rating '4' (Lancar) pada kosakata '行く'",
            "expected": "Data terkirim via POST /api/v1/srs/review, interval berikutnya dihitung ulang berbasis algoritma SM-2, kartu review beralih.",
            "actual": "Berhasil sesuai harapan, interval ter-update di database.", "status": "[√] Valid"
        },
        {
            "no": "3", "skenario": "Menilai kartu yang lupa",
            "case": "1. Klik Tampilkan Jawaban\n2. Klik tombol rating '0' (Lupa)",
            "data": "Memilih rating '0' (Lupa)",
            "expected": "Repetisi kartu direset ke 0, interval diatur ke 1 hari, kartu diulang dalam sesi review hari ini.",
            "actual": "Berhasil sesuai harapan, kartu dimasukkan kembali ke antrean terdekat.", "status": "[√] Valid"
        }
    ]
    render_table("Tabel 4.8 Testing Halaman SRS Review", srs_data)

    # 1.9 Homepage Speaking Mode
    add_heading_3("1.9 Testing Homepage Speaking Mode")
    add_paragraph("Pengujian speaking mode dilakukan untuk menguji tangkapan audio mic (Speech-to-Text) serta penilaian akurasi pelafalan dari backend.")
    speaking_data = [
        {
            "no": "1", "skenario": "Perekaman input lisan pengguna",
            "case": "1. Klik dan tahan tombol mic\n2. Ucapkan kalimat Jepang\n3. Lepas tombol mic",
            "data": "Email: testing@gmail.com\nKalimat diucapkan: 'こんにちは'",
            "expected": "Cincin visual gelombang suara aktif bergerak, dan ucapan berhasil ditranskripsikan ke teks Jepang di layar via STT.",
            "actual": "Berhasil sesuai harapan, transkrip lisan muncul di layar.", "status": "[√] Valid"
        },
        {
            "no": "2", "skenario": "Penilaian akurasi pengucapan lisan",
            "case": "1. Kirim hasil ucapan\n2. Amati bubble chat user",
            "data": "Ucapan: 'これはペンです'",
            "expected": "Bubble chat user menampilkan badge persentase akurasi (misal: 'Akurasi 95%'). Klik badge menampilkan detail akurasi fonetis.",
            "actual": "Berhasil sesuai harapan, badge akurasi berwarna hijau muncul.", "status": "[√] Valid"
        }
    ]
    render_table("Tabel 4.9 Testing Homepage Speaking Mode", speaking_data)

    # 1.10 Halaman Profile
    add_heading_3("1.10 Testing Halaman Profile")
    add_paragraph("Pengujian halaman profile dilakukan untuk memverifikasi pemuatan grafik radar kemahiran, streak harian, dan ekspor riwayat progres belajar.")
    profile_data = [
        {
            "no": "1", "skenario": "Memuat radar chart 5 kompetensi",
            "case": "1. Klik tab Profile\n2. Amati grafik radar kemahiran BKT",
            "data": "Email: Raincallin@gmail.com",
            "expected": "Svelte component RadarChart merender grafik radar 5 dimensi secara dinamis sesuai status kognitif BKT user.",
            "actual": "Berhasil sesuai harapan, radar chart tergambar rapi.", "status": "[√] Valid"
        },
        {
            "no": "2", "skenario": "Mengubah target belajar harian",
            "case": "1. Klik tombol pensil (edit target)\n2. Masukkan target baru\n3. Klik Simpan Target",
            "data": "Vocab: 15, Grammar: 3, Review: 8, Study: 20 menit",
            "expected": "API PUT dipanggil, target tersimpan di Supabase database, UI progress bar ter-update.",
            "actual": "Berhasil sesuai harapan, target baru berhasil disimpan.", "status": "[√] Valid"
        },
        {
            "no": "3", "skenario": "Memvalidasi streak calendar",
            "case": "1. Amati grid kalender 90 hari terakhir",
            "data": "Data log aktivitas belajar",
            "expected": "Kalender heatmap 90 hari dirender dengan gradasi warna hijau sesuai intensitas aktivitas harian.",
            "actual": "Berhasil sesuai harapan, heatmap terisi kotak hijau.", "status": "[√] Valid"
        }
    ]
    render_table("Tabel 4.10 Testing Halaman Profile", profile_data)

    # 1.11 Halaman Achievement
    add_heading_3("1.11 Testing Halaman Achievement")
    add_paragraph("Pengujian halaman achievement dilakukan untuk memastikan perolehan lencana penghargaan (badges) dan visualisasi pangkat samurai berdasarkan XP.")
    achievement_data = [
        {
            "no": "1", "skenario": "Klaim lencana dan pangkat samurai",
            "case": "1. Klik tab Achievement\n2. Amati badges dan tingkat pangkat",
            "data": "Email: Satyakrisna89=0@gmail.com\nLevel user: 4",
            "expected": "Sistem menampilkan lencana diperoleh menyala berwarna, lencana terkunci grayscale, dan pangkat samurai dinamis (Ronin) berdasarkan level.",
            "actual": "Berhasil sesuai harapan, pangkat Ronin dan badge aktif menyala.", "status": "[√] Valid"
        },
        {
            "no": "2", "skenario": "Melihat jalur penguasaan materi",
            "case": "1. Klik tab sub-menu 'Materi' pada halaman pencapaian",
            "data": "State mastery_path",
            "expected": "Menampilkan representasi grafis simpul materi kurikulum dengan indikator warna hijau (MASTERED) atau merah/kuning (STRUGGLING).",
            "actual": "Berhasil sesuai harapan, grafik simpul materi tergambar dinamis.", "status": "[√] Valid"
        }
    ]
    render_table("Tabel 4.11 Testing Halaman Achievement", achievement_data)


    # 2. White-box testing
    add_heading_2("2. White-box testing")
    add_paragraph(
        "Pengujian white-box (kotak putih) dilakukan untuk menguji kebenaran alur logika internal, percabangan, "
        "kondisi batas, serta cakupan jalur eksekusi struktural dari modul-modul penting pada sistem TVJP A.L.I.S.A. "
        "Metode pengujian struktural yang digunakan adalah Basis Path Testing, di mana setiap method/event handler "
        "dimodelkan menjadi grafik aliran kontrol (Control Flow Graph atau CFG), dihitung kompleksitas siklomatisnya "
        "untuk mengidentifikasi jumlah jalur independen, didefinisikan rincian jalur logikanya, serta divalidasi "
        "menggunakan skenario unit testing otomatis."
    )

    # 2.1 Pengujian whitebox pada method submitReview()
    add_heading_3("2.1 Pengujian whitebox pada method submitReview()")
    add_paragraph(
        "Method submitReview(quality) terletak pada komponen SRSReview.svelte. Method ini bertugas untuk mencatat "
        "ulasan kartu memori (spaced repetition system) dari pengguna, menghitung perolehan poin pengalaman (XP) "
        "secara dinamis, mengirimkan data log review ke server backend via API POST /api/v1/srs/review, serta memajukan "
        "indeks kartu ke kartu berikutnya atau membersihkan antrean ulasan jika telah selesai."
    )
    
    add_heading_3("2.1.1 Source code yang diuji")
    add_paragraph(
        "Berikut adalah potongan kode logika utama pada method submitReview() yang telah diberi nomor simpul (node) "
        "untuk keperluan analisis grafik alur kendali (Control Flow Graph):"
    )
    
    code_p = doc.add_paragraph()
    code_p.paragraph_format.left_indent = Inches(0.2)
    code_run = code_p.add_run(
        "async function submitReview(quality) {\n"
        "    // [Node 1] Ambil item kartu aktif\n"
        "    const item = dueItems[currentIndex];\n"
        "    // [Node 2] Cek jika item kosong atau pengguna belum login\n"
        "    if (!item || !$user) return; // [Node 2a] Return early\n\n"
        "    // [Node 3] Hitung XP dan ulasan\n"
        "    const gain = quality >= 3 ? 5 : 2;\n"
        "    xpEarned += gain;\n"
        "    reviewsCompleted++;\n\n"
        "    // [Node 4] POST API review\n"
        "    try {\n"
        "        await fetch(\"http://localhost:8000/api/v1/srs/review\", {\n"
        "            method: \"POST\",\n"
        "            headers: { \"Content-Type\": \"application/json\" },\n"
        "            body: JSON.stringify({ ... })\n"
        "        });\n"
        "    } catch(e) { ... }\n\n"
        "    // [Node 5] Cek kartu sisa\n"
        "    if (currentIndex < dueItems.length - 1) {\n"
        "        // [Node 5a] Majukan indeks\n"
        "        currentIndex++;\n"
        "        await loadCurrentNodeDetails();\n"
        "    } else {\n"
        "        // [Node 5b] Bersihkan antrean\n"
        "        dueItems = [];\n"
        "    }\n"
        "    // [Node 6] Exit\n"
        "}"
    )
    code_run.font.name = "Courier New"
    code_run.font.size = Pt(8.5)
    
    rows_srs_struct = [
        {
            "no": "1", "skenario": "Pengguna belum login atau antrean kosong",
            "case": "Simulasi method dipanggil dengan user = null atau dueItems = []",
            "data": "user = null, dueItems = []",
            "expected": "Sistem keluar dari method melalui early return seketika.",
            "actual": "Berhasil sesuai harapan, early return sukses.", "status": "[√] Valid"
        },
        {
            "no": "2", "skenario": "Kartu maju dalam antrean",
            "case": "Simulasi ulasan dikirim dan masih ada kartu tersisa",
            "data": "dueItems = [item1, item2], quality = 4",
            "expected": "XP bertambah 5, currentIndex menjadi 1, API POST terpanggil.",
            "actual": "Berhasil sesuai harapan, currentIndex bertambah.", "status": "[√] Valid"
        },
        {
            "no": "3", "skenario": "Ulasan pada kartu terakhir selesai",
            "case": "Simulasi ulasan dikirim pada kartu terakhir",
            "data": "dueItems = [item1], quality = 2",
            "expected": "XP bertambah 2, dueItems dikosongkan, API POST terpanggil.",
            "actual": "Berhasil sesuai harapan, dueItems menjadi kosong.", "status": "[√] Valid"
        }
    ]
    render_table("Tabel 4.12 Skenario Pengujian Struktural submitReview()", rows_srs_struct)
    
    add_heading_3("2.1.2 Control flow Graph(CFG)")
    add_paragraph(
        "Grafik aliran kontrol (Control Flow Graph) dari method submitReview() dimodelkan untuk mendeteksi "
        "seluruh jalur logika percabangan. Kode XML mxGraph (draw.io) berikut dapat digunakan untuk merepresentasikan "
        "grafik aliran kontrol"
    )
    
    xml_p = doc.add_paragraph()
    xml_p.paragraph_format.left_indent = Inches(0.2)
    xml_run = xml_p.add_run(r'''<mxfile host="Electron" modified="2026-07-12T10:00:00.000Z" agent="5.0" version="20.0.0" type="device">
  <diagram id="cfg-submit-review" name="CFG submitReview">
    <mxGraphModel dx="1000" dy="1000" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="827" pageHeight="1169" math="0" shadow="0">
      <root>
        <mxCell id="0" />
        <mxCell id="1" parent="0" />
        <mxCell id="n1" value="Node 1: Entry" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="40" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n2" value="Node 2: if (!item || !$user)" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="160" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n2a" value="Node 2a: return" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="240" y="160" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n3" value="Node 3: Scoring &amp; XP" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="280" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n4" value="Node 4: POST API log" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="400" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n5" value="Node 5: if (currentIndex < length - 1)" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="520" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n5a" value="Node 5a: currentIndex++" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="280" y="640" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n5b" value="Node 5b: dueItems = []" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="480" y="640" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n6" value="Node 6: Exit" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="760" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="e1" parent="1" source="n1" target="n2" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e2" value="Ya" parent="1" source="n2" target="n2a" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e3" value="Tidak" parent="1" source="n2" target="n3" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e4" parent="1" source="n3" target="n4" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e5" parent="1" source="n4" target="n5" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e6" value="Ya" parent="1" source="n5" target="n5a" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e7" value="Tidak" parent="1" source="n5" target="n5b" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e8" parent="1" source="n2a" target="n6" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e9" parent="1" source="n5a" target="n6" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e10" parent="1" source="n5b" target="n6" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
      </root>
    </mxGraphModel>
  </diagram>
</mxfile>''')
    xml_run.font.name = "Courier New"
    xml_run.font.size = Pt(7.5)
    
    add_heading_3("2.1.3 Cyclomatic complexity")
    add_paragraph("Rumus perhitungan Cyclomatic Complexity berdasarkan teori graf adalah sebagai berikut:")
    
    math_p = doc.add_paragraph()
    math_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math_run = math_p.add_run("V(G) = E - N + 2")
    math_run.bold = True
    math_run.font.name = "Times New Roman"
    math_run.font.size = Pt(11)
    
    add_paragraph("Di mana:\n- E = 10 (jumlah panah alur kendali pada CFG)\n- N = 9 (jumlah simpul proses/keputusan pada CFG)")
    add_paragraph("Maka perhitungan nilainya adalah:\n  V(G) = 10 - 9 + 2 = 3")
    add_paragraph("Hasil perhitungan menunjukkan nilai Cyclomatic Complexity sebesar 3, yang menyatakan terdapat 3 jalur logika independen (independent path) yang wajib diuji.")
    
    add_heading_3("2.1.4 Independent path")
    add_paragraph("Berdasarkan perhitungan kompleksitas, rincian dari 3 jalur logika independen (independent path) yang diuji adalah:")
    add_paragraph("Path 1 (1 -> 2 -> 2a -> 6): Jalur di mana method langsung keluar (early return) karena kondisi user atau kartu kosong.")
    add_paragraph("Path 2 (1 -> 2 -> 3 -> 4 -> 5 -> 5a -> 6): Jalur ulasan dikirim sukses untuk kartu non-terakhir, yang memajukan indeks kartu (currentIndex++).")
    add_paragraph("Path 3 (1 -> 2 -> 3 -> 4 -> 5 -> 5b -> 6): Jalur ulasan dikirim sukses untuk kartu terakhir, yang membersihkan antrean kartu (dueItems = []).")
    
    add_heading_3("2.1.5 Implementasi dan Hasil Unit testing")
    add_paragraph(
        "Pengujian white-box dengan teknik basis path testing dilakukan menggunakan unit testing otomatis unittest pada "
        "backend dengan memodelkan logic handler Svelte secara presisi. Setiap jalur independen diuji dengan menyimulasikan "
        "masukan state yang berbeda untuk memastikan seluruh percabangan berjalan sesuai spesifikasi melalui perintah "
        "python run_whitebox_submitReview.py."
    )
    
    rows_srs_unit = [
        {
            "path": "1", "skenario": "Ulasan diabaikan karena user atau item kosong",
            "input": "user = None, dueItems = []",
            "expected": "Sistem keluar dari method melalui early return seketika",
            "status": "PASS"
        },
        {
            "path": "2", "skenario": "Perekaman sukses dan kartu maju",
            "input": "user = {'id': 'usr_1'}, dueItems = [item1, item2], quality = 4",
            "expected": "XP bertambah 5, currentIndex menjadi 1, status advanced",
            "status": "PASS"
        },
        {
            "path": "3", "skenario": "Perekaman sukses pada kartu terakhir",
            "input": "user = {'id': 'usr_1'}, dueItems = [item1], quality = 2",
            "expected": "XP bertambah 2, dueItems dikosongkan ([]), status finished",
            "status": "PASS"
        }
    ]
    render_table_5("Tabel 4.13 Hasil Unit Testing Method submitReview()", rows_srs_unit)
    
    add_paragraph(
        "Seluruh 3 jalur independen berhasil dieksekusi dan menghasilkan output sesuai ekspektasi. Hal ini menunjukkan "
        "bahwa validasi pada method submitReview() berjalan dengan baik dalam mencegah data tidak valid tersimpan ke database. "
        "Selain itu, integrasi pengiriman poin pengalaman (XP) dan mutasi status antrean kartu memori berjalan secara "
        "atomik sesuai spesifikasi yang diharapkan, sebagaimana ditunjukkan pada Gambar 4.120 berikut."
    )
    
    add_paragraph(
        "*(Gambar 4.120 Hasil Unit Testing Method submitReview() ditunjukkan pada tangkapan layar terminal)*",
        italic=True
    )
    
    add_heading_3("2.1.6 Kesimpulan")
    add_paragraph(
        "Method submitReview() memiliki nilai Cyclomatic Complexity sebesar 3. Seluruh 3 jalur independen berhasil "
        "dieksekusi dengan sukses (PASS) sesuai dengan spesifikasi rancangan logika sistem."
    )



    add_heading_3("2.2 Pengujian whitebox pada method handleAnswer()")
    add_paragraph(
        "Method handleAnswer(userAnswer) terletak pada komponen QuestEngine.svelte. Method ini bertugas untuk menangani "
        "proses pengiriman jawaban kuis oleh pengguna, memvalidasi apakah sistem sedang dalam proses evaluasi (mencegah klik ganda), "
        "menilai kebenaran jawaban, menghitung poin skor secara dinamis (berdasarkan jumlah kesalahan pengerjaan dan penggunaan "
        "bantuan petunjuk/hint), memperbarui kemajuan belajar (mastery status) ke backend server jika pengguna telah masuk log "
        "(logged in), serta memajukan kuis ke pertanyaan berikutnya."
    )

    add_heading_3("2.2.1 Source code yang diuji")
    add_paragraph(
        "Berikut adalah potongan kode logika utama pada method handleAnswer() yang telah diberi nomor simpul (node) "
        "untuk keperluan analisis grafik alur kendali (Control Flow Graph):"
    )

    code_p2 = doc.add_paragraph()
    code_p2.paragraph_format.left_indent = Inches(0.2)
    code_run2 = code_p2.add_run(
        "async function handleAnswer(userAnswer) {\n"
        "    // [Node 1] Entry\n"
        "    if (isEvaluating || showAiFeedback) return; // [Node 2, 2a]\n"
        "    isEvaluating = true;\n"
        "    const isCorrect = checkAnswer(currentQuestion, userAnswer); // [Node 3]\n\n"
        "    if (isCorrect) { // [Node 4]\n"
        "        // [Node 5] Jawaban Benar\n"
        "        sfx.play(\"success\");\n"
        "        score += earnedScore;\n"
        "        if ($user && currentQuestion.node_id) { // [Node 6]\n"
        "            // [Node 6a] Log Mastery\n"
        "            await logMastery(masteryStatus);\n"
        "        }\n"
        "        setTimeout(() => { ... }, 900); // [Node 7]\n"
        "    } else {\n"
        "        // [Node 8] Jawaban Salah\n"
        "        sfx.play(\"failure\");\n"
        "        wrongAttempts[currentQuestion.id]++;\n"
        "        isEvaluating = false;\n"
        "    }\n"
        "    // [Node 9] Exit\n"
        "}"
    )
    code_run2.font.name = "Courier New"
    code_run2.font.size = Pt(8.5)

    rows_ans_struct = [
        {
            "no": "1", "skenario": "Mencegah pengerjaan ganda saat evaluasi berjalan",
            "case": "Simulasi double click pada tombol jawaban",
            "data": "isEvaluating = true",
            "expected": "Method langsung dihentikan melalui early return seketika.",
            "actual": "Berhasil sesuai harapan, early return sukses.", "status": "[√] Valid"
        },
        {
            "no": "2", "skenario": "Jawaban benar pertama tanpa hint",
            "case": "Jawaban benar langsung dengan user login",
            "data": "isCorrect = true, attempts = 0, wasHinted = false, user = active",
            "expected": "Skor +10, API log-mastery terpanggil, timer kuis terpicu.",
            "actual": "Berhasil sesuai harapan, skor bertambah 10.", "status": "[√] Valid"
        },
        {
            "no": "3", "skenario": "Jawaban benar kedua dengan hint (guest)",
            "case": "Jawaban benar lanjutan dengan petunjuk tanpa login",
            "data": "isCorrect = true, attempts = 1, wasHinted = true, user = null",
            "expected": "Skor +4, API log-mastery dilewati, timer kuis terpicu.",
            "actual": "Berhasil sesuai harapan, skor bertambah 4.", "status": "[√] Valid"
        },
        {
            "no": "4", "skenario": "Jawaban salah pada pengerjaan",
            "case": "Percobaan jawaban salah",
            "data": "isCorrect = false",
            "expected": "Attempts bertambah 1, lastAnswerWrong menjadi true, skor tetap.",
            "actual": "Berhasil sesuai harapan, attempts bertambah.", "status": "[√] Valid"
        }
    ]
    render_table("Tabel 4.14 Skenario Pengujian Struktural handleAnswer()", rows_ans_struct)

    add_heading_3("2.2.2 Control flow Graph(CFG)")
    add_paragraph(
        "Grafik aliran kontrol (Control Flow Graph) dari method handleAnswer() dimodelkan untuk mendeteksi "
        "seluruh jalur logika percabangan. Kode XML mxGraph (draw.io) berikut merepresentasikan CFG handleAnswer():"
    )

    xml_p2 = doc.add_paragraph()
    xml_p2.paragraph_format.left_indent = Inches(0.2)
    xml_run2 = xml_p2.add_run(r'''<mxfile host="Electron" modified="2026-07-12T10:00:00.000Z" agent="5.0" version="20.0.0" type="device">
  <diagram id="cfg-handle-answer" name="CFG handleAnswer">
    <mxGraphModel dx="1000" dy="1000" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="827" pageHeight="1169" math="0" shadow="0">
      <root>
        <mxCell id="0" />
        <mxCell id="1" parent="0" />
        <mxCell id="n1" value="Node 1: Entry" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="40" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n2" value="Node 2: if (isEvaluating || showAiFeedback)" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="160" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n2a" value="Node 2a: return" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="240" y="160" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n3" value="Node 3: isEvaluating = true &amp; check" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="280" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n4" value="Node 4: if (isCorrect)" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="400" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n5" value="Node 5: Score Calc" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="280" y="520" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n6" value="Node 6: if ($user &amp;&amp; node_id)" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="280" y="640" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n6a" value="Node 6a: fetch Log" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="180" y="760" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n7" value="Node 7: setTimeout" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="280" y="880" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n8" value="Node 8: Attempts++" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="480" y="520" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n9" value="Node 9: Exit" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="1000" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="e1" parent="1" source="n1" target="n2" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e2" value="Ya" parent="1" source="n2" target="n2a" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e3" value="Tidak" parent="1" source="n2" target="n3" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e4" parent="1" source="n3" target="n4" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e5" value="Ya" parent="1" source="n4" target="n5" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e6" value="Tidak" parent="1" source="n4" target="n8" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e7" parent="1" source="n5" target="n6" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e8" value="Ya" parent="1" source="n6" target="n6a" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e9" value="Tidak" parent="1" source="n6" target="n7" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e10" parent="1" source="n6a" target="n7" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e11" parent="1" source="n7" target="n9" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e12" parent="1" source="n8" target="n9" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e13" parent="1" source="n2a" target="n9" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
      </root>
    </mxGraphModel>
  </diagram>
</mxfile>''')
    xml_run2.font.name = "Courier New"
    xml_run2.font.size = Pt(7.5)

    add_heading_3("2.2.3 Cyclomatic complexity")
    add_paragraph("Rumus perhitungan Cyclomatic Complexity berdasarkan teori graf adalah sebagai berikut:")

    math_p2 = doc.add_paragraph()
    math_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math_run2 = math_p2.add_run("V(G) = P + 1")
    math_run2.bold = True
    math_run2.font.name = "Times New Roman"
    math_run2.font.size = Pt(11)

    add_paragraph("Di mana:\n- P = 3 (jumlah simpul keputusan/predicate node pada alur kendali, yaitu Node 2, Node 4, dan Node 6)")
    add_paragraph("Maka perhitungan nilainya adalah:\n  V(G) = 3 + 1 = 4")
    add_paragraph("Hasil perhitungan menunjukkan nilai Cyclomatic Complexity sebesar 4, menyatakan terdapat 4 jalur logika independen (independent path) yang wajib diuji.")

    add_heading_3("2.2.4 Independent path")
    add_paragraph("Berdasarkan perhitungan kompleksitas, rincian dari 4 jalur logika independen (independent path) yang diuji adalah:")
    add_paragraph("Path 1 (1 -> 2 -> 2a -> 9): Aliran eksekusi di mana sistem langsung keluar (early return) karena sedang mengevaluasi.")
    add_paragraph("Path 2 (1 -> 2 -> 3 -> 4 -> 5 -> 6 -> 6a -> 7 -> 9): Aliran jawaban benar pada percobaan pertama tanpa hint, mengirim log mastery ke server.")
    add_paragraph("Path 3 (1 -> 2 -> 3 -> 4 -> 5 -> 6 -> 7 -> 9): Aliran jawaban benar pada percobaan kedua dengan hint oleh pengguna tamu.")
    add_paragraph("Path 4 (1 -> 2 -> 3 -> 4 -> 8 -> 9): Aliran jawaban salah dari pengguna, meningkatkan attempts.")

    add_heading_3("2.2.5 Implementasi dan Hasil Unit testing")
    add_paragraph(
        "Pengujian white-box dengan teknik basis path testing dilakukan menggunakan unit testing otomatis unittest pada "
        "backend dengan memodelkan logic handler Svelte secara presisi. Setiap jalur independen diuji dengan menyimulasikan "
        "masukan state yang berbeda untuk memastikan seluruh percabangan berjalan sesuai spesifikasi melalui perintah "
        "python run_whitebox_handleAnswer.py."
    )

    rows_ans_unit = [
        {
            "path": "1", "skenario": "Mencegah pengerjaan ganda saat evaluasi berjalan",
            "input": "isEvaluating = true",
            "expected": "Sistem keluar dari method melalui early return seketika",
            "status": "PASS"
        },
        {
            "path": "2", "skenario": "Jawaban benar pada percobaan pertama tanpa hint",
            "input": "isCorrect = true, attempts = 0, wasHinted = false, user = active",
            "expected": "Poin skor bertambah 10, API log-mastery terpanggil dengan status \"MASTERED\"",
            "status": "PASS"
        },
        {
            "path": "3", "skenario": "Jawaban benar pada percobaan kedua dengan hint (tamu)",
            "input": "isCorrect = true, attempts = 1, wasHinted = true, user = null",
            "expected": "Poin skor bertambah 4, API log-mastery tidak terpanggil",
            "status": "PASS"
        },
        {
            "path": "4", "skenario": "Jawaban salah pada pengerjaan",
            "input": "isCorrect = false",
            "expected": "wrongAttempts bertambah 1, skor tetap, lastAnswerWrong = true",
            "status": "PASS"
        }
    ]
    render_table_5("Tabel 4.15 Hasil Pengujian Unit Testing Method handleAnswer()", rows_ans_unit)

    add_paragraph(
        "Seluruh 4 jalur independen berhasil dieksekusi dan menghasilkan output sesuai ekspektasi. Hal ini menunjukkan "
        "bahwa validasi pada method handleAnswer() berjalan dengan baik dalam mencegah data tidak valid tersimpan ke database. "
        "Selain itu, integrasi pengisian skor adaptif dan pencatatan riwayat kemajuan siswa berjalan secara aman sesuai "
        "spesifikasi yang diharapkan, sebagaimana ditunjukkan pada Gambar 4.121 berikut."
    )

    add_paragraph(
        "*(Gambar 4.121 Hasil Unit Testing Method handleAnswer() ditunjukkan pada tangkapan layar terminal)*",
        italic=True
    )

    add_heading_3("2.2.6 Kesimpulan")
    add_paragraph(
        "Method handleAnswer() memiliki nilai Cyclomatic Complexity sebesar 4. Seluruh 4 jalur independen berhasil "
        "dieksekusi dengan sukses (PASS) sesuai dengan spesifikasi rancangan logika sistem."
    )

    add_heading_3("2.3 Pengujian whitebox pada method finishExam()")
    add_paragraph(
        "Method finishExam(timeUp) terletak pada komponen ExamEngine.svelte. Method ini bertugas untuk menghentikan "
        "sesi ujian yang sedang berjalan (baik karena waktu habis timeUp maupun karena pengguna menyelesaikan secara manual), "
        "menghentikan interval pewaktu (timer), mengirimkan perolehan skor ujian pengguna ke backend server via API POST "
        "/api/v1/exam/submit jika pengguna telah masuk log (logged in), menangani error koneksi database, serta mengarahkan "
        "tampilan halaman ke tab hasil (result)."
    )

    add_heading_3("2.3.1 Source code yang diuji")
    add_paragraph(
        "Berikut adalah potongan kode logika utama pada method finishExam() yang telah diberi nomor simpul (node) "
        "untuk keperluan analisis grafik alur kendali (Control Flow Graph):"
    )

    code_p3 = doc.add_paragraph()
    code_p3.paragraph_format.left_indent = Inches(0.2)
    code_run3 = code_p3.add_run(
        "async function finishExam(timeUp = false) {\n"
        "    // [Node 1] Entry\n"
        "    if (examFinished) return; // [Node 2, 2a]\n"
        "    examFinished = true;\n"
        "    if (timerInterval) clearInterval(timerInterval); // [Node 3]\n\n"
        "    if ($user) { // [Node 4]\n"
        "        try { // [Node 5] Kirim data skor\n"
        "            const response = await fetch(\"http://localhost:8000/api/v1/exam/submit\", {\n"
        "                method: \"POST\",\n"
        "                headers: { \"Content-Type\": \"application/json\" },\n"
        "                body: JSON.stringify({ ... })\n"
        "            });\n"
        "            if (!response.ok) throw new Error(\"...\"); // [Node 6, 6a]\n"
        "            const resData = await response.json();\n"
        "            examResultId = resData.result_id; // [Node 6b]\n"
        "        } catch(e) {\n"
        "            saveError = \"...\"; // [Node 6a] Tangkap error\n"
        "        }\n"
        "    }\n"
        "    activeTab = \"result\"; // [Node 8]\n"
        "    // [Node 9] Exit\n"
        "}"
    )
    code_run3.font.name = "Courier New"
    code_run3.font.size = Pt(8.5)

    rows_ex_struct = [
        {
            "no": "1", "skenario": "Ujian diselesaikan oleh pengguna tamu (Demo/Guest)",
            "case": "Ujian selesai tanpa login",
            "data": "user = null, timeUp = False",
            "expected": "examFinished bernilai true, API dilewati, activeTab beralih ke 'result'.",
            "actual": "Berhasil sesuai harapan, tab beralih.", "status": "[√] Valid"
        },
        {
            "no": "2", "skenario": "Ujian selesai otomatis (waktu habis) oleh user login",
            "case": "Ujian waktu habis dengan user login",
            "data": "user = {'id': 'usr_1'}, timeUp = True",
            "expected": "examFinished bernilai true, API POST sukses terkirim, examResultId tersimpan.",
            "actual": "Berhasil sesuai harapan, API sukses.", "status": "[√] Valid"
        },
        {
            "no": "3", "skenario": "Ujian selesai oleh user login namun terjadi kegagalan server",
            "case": "API submit return 500 error",
            "data": "user = {'id': 'usr_1'}, timeUp = False",
            "expected": "Ujian selesai, saveError diset dengan aman, activeTab beralih ke 'result'.",
            "actual": "Berhasil sesuai harapan, error ditangkap.", "status": "[√] Valid"
        }
    ]
    render_table("Tabel 4.16 Skenario Pengujian Struktural finishExam()", rows_ex_struct)

    add_heading_3("2.3.2 Control flow Graph(CFG)")
    add_paragraph(
        "Grafik aliran kontrol (Control Flow Graph) dari method finishExam() dimodelkan untuk mendeteksi "
        "seluruh alur logika percabangan. Kode XML mxGraph (draw.io) berikut merepresentasikan CFG finishExam():"
    )

    xml_p3 = doc.add_paragraph()
    xml_p3.paragraph_format.left_indent = Inches(0.2)
    xml_run3 = xml_p3.add_run(r'''<mxfile host="Electron" modified="2026-07-12T10:00:00.000Z" agent="5.0" version="20.0.0" type="device">
  <diagram id="cfg-finish-exam" name="CFG finishExam">
    <mxGraphModel dx="1000" dy="1000" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="827" pageHeight="1169" math="0" shadow="0">
      <root>
        <mxCell id="0" />
        <mxCell id="1" parent="0" />
        <mxCell id="n1" value="Node 1: Entry" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="40" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n2" value="Node 2: if (examFinished)" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="160" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n2a" value="Node 2a: return" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="240" y="160" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n3" value="Node 3: examFinished = true &amp; Clear" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="280" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n4" value="Node 4: if ($user)" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="400" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n5" value="Node 5: fetch POST submit" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="520" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n6" value="Node 6: if (!response.ok)" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="640" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n6a" value="Node 6a: catch / saveError" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="280" y="760" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n6b" value="Node 6b: examResultId =" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="480" y="760" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n8" value="Node 8: activeTab = 'result'" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="880" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n9" value="Node 9: Exit" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="1000" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="e1" parent="1" source="n1" target="n2" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e2" value="Ya" parent="1" source="n2" target="n2a" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e3" value="Tidak" parent="1" source="n2" target="n3" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e4" parent="1" source="n3" target="n4" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e5" value="Ya" parent="1" source="n4" target="n5" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e6" value="Tidak" parent="1" source="n4" target="n8" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e7" parent="1" source="n5" target="n6" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e8" value="Ya" parent="1" source="n6" target="n6a" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e9" value="Tidak" parent="1" source="n6" target="n6b" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e10" parent="1" source="n6a" target="n8" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e11" parent="1" source="n6b" target="n8" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e12" parent="1" source="n8" target="n9" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e13" parent="1" source="n2a" target="n9" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
      </root>
    </mxGraphModel>
  </diagram>
</mxfile>''')
    xml_run3.font.name = "Courier New"
    xml_run3.font.size = Pt(7.5)

    add_heading_3("2.3.3 Cyclomatic complexity")
    add_paragraph("Rumus perhitungan Cyclomatic Complexity berdasarkan teori graf adalah sebagai berikut:")

    math_p3 = doc.add_paragraph()
    math_p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math_run3 = math_p3.add_run("V(G) = P + 1")
    math_run3.bold = True
    math_run3.font.name = "Times New Roman"
    math_run3.font.size = Pt(11)

    add_paragraph("Di mana:\n- P = 2 (jumlah simpul keputusan/predicate node pada alur kendali, yaitu Node 4 dan Node 6)")
    add_paragraph("Maka perhitungan nilainya adalah:\n  V(G) = 2 + 1 = 3")
    add_paragraph("Hasil perhitungan menunjukkan nilai Cyclomatic Complexity sebesar 3, menyatakan terdapat 3 jalur logika independen (independent path) utama yang wajib diuji.")

    add_heading_3("2.3.4 Independent path")
    add_paragraph("Berdasarkan perhitungan kompleksitas, rincian dari 3 jalur logika independen (independent path) yang diuji adalah:")
    add_paragraph("Path 1 (1 -> 2 -> 3 -> 4 -> 8 -> 9): Aliran eksekusi di mana pengguna tamu menyelesaikan ujian tanpa login.")
    add_paragraph("Path 2 (1 -> 2 -> 3 -> 4 -> 5 -> 6 -> 6b -> 8 -> 9): Aliran eksekusi di mana user login menyelesaikan ujian, data terkirim sukses via API POST.")
    add_paragraph("Path 3 (1 -> 2 -> 3 -> 4 -> 5 -> 6 -> 6a -> 8 -> 9): Aliran eksekusi di mana user login menyelesaikan ujian namun API POST gagal.")

    add_heading_3("2.3.5 Implementasi dan Hasil Unit testing")
    add_paragraph(
        "Pengujian white-box dengan teknik basis path testing dilakukan menggunakan unit testing otomatis unittest pada "
        "backend dengan memodelkan logic handler Svelte secara presisi. Setiap jalur independen diuji dengan menyimulasikan "
        "masukan state yang berbeda untuk memastikan seluruh percabangan berjalan sesuai spesifikasi melalui perintah "
        "python run_whitebox_finishExam.py."
    )

    rows_ex_unit = [
        {
            "path": "1", "skenario": "Ujian diselesaikan oleh pengguna tamu (Demo/Guest)",
            "input": "user = None, timeUp = False",
            "expected": "examFinished menjadi true, API dilewati, activeTab beralih ke \"result\"",
            "status": "PASS"
        },
        {
            "path": "2", "skenario": "Ujian selesai otomatis (waktu habis) oleh user login",
            "input": "user = {'id': 'usr_1'}, timeUp = True",
            "expected": "examFinished menjadi true, API POST sukses terkirim, examResultId tersimpan",
            "status": "PASS"
        },
        {
            "path": "3", "skenario": "Ujian selesai oleh user login namun terjadi kegagalan server",
            "input": "user = {'id': 'usr_1'}, timeUp = False, API = Failure",
            "expected": "Ujian tetap selesai, saveError diset aman, activeTab beralih ke \"result\"",
            "status": "PASS"
        }
    ]
    render_table_5("Tabel 4.16 Hasil Pengujian Unit Testing Method finishExam()", rows_ex_unit)

    add_paragraph(
        "Seluruh 3 jalur independen berhasil dieksekusi dan menghasilkan output sesuai ekspektasi. Hal ini menunjukkan "
        "bahwa penanganan logika penyelesaian sesi ujian, perlindungan waktu habis (timeout), dan penanganan kegagalan API "
        "pada method finishExam() berjalan dengan aman dari risiko kegagalan fatal. Detail hasil eksekusi ditunjukkan "
        "pada Gambar 4.122 berikut."
    )

    add_paragraph(
        "*(Gambar 4.122 Hasil Unit Testing Method finishExam() ditunjukkan pada tangkapan layar terminal)*",
        italic=True
    )

    add_heading_3("2.3.6 Kesimpulan")
    add_paragraph(
        "Method finishExam() memiliki nilai Cyclomatic Complexity sebesar 3. Seluruh 3 jalur independen berhasil "
        "dieksekusi dengan sukses (PASS) sesuai dengan spesifikasi rancangan logika sistem."
    )

    add_heading_3("2.4 Pengujian whitebox pada method sendChat()")
    add_paragraph(
        "Method sendChat(text) terletak pada komponen utama routes/+page.svelte. Method ini bertugas untuk menangani "
        "proses pengiriman pesan percakapan oleh pengguna ke chatbot edukasi berbasis WebSocket, melakukan pembersihan dan "
        "validasi input (mencegah pesan kosong atau pengiriman berulang saat status pemrosesan loading aktif), "
        "menginisialisasi parameter visualisasi animasi thinking, serta mengoordinasikan pengiriman data ke server "
        "backend melalui protokol WebSocket."
    )

    add_heading_3("2.4.1 Source code yang diuji")
    add_paragraph(
        "Berikut adalah potongan kode logika utama pada method sendChat() yang telah diberi nomor simpul (node) "
        "untuk keperluan analisis grafik alur kendali (Control Flow Graph):"
    )

    code_p4 = doc.add_paragraph()
    code_p4.paragraph_format.left_indent = Inches(0.2)
    code_run4 = code_p4.add_run(
        "async function sendChat(text = null) {\n"
        "    // [Node 1] Entry\n"
        "    const userText = (text || query).trim();\n"
        "    if (!userText || loading) return; // [Node 2, 2a]\n"
        "    query = \"\";\n\n"
        "    // [Node 3] Update state chatStore, loading, & thinking\n"
        "    chatStore.update((s) => ({ ... }));\n"
        "    loading = true;\n"
        "    startThinkingAnimation();\n\n"
        "    // [Node 4] Ambil WebSocket & cek status\n"
        "    const ws = getWS();\n"
        "    if (ws.readyState === WebSocket.OPEN) {\n"
        "        dispatch(); // [Node 5, 5a]\n"
        "    } else {\n"
        "        // [Node 6] Connecting listener\n"
        "        ws.addEventListener(\"open\", () => dispatch(), { once: true });\n"
        "    }\n"
        "    // [Node 7] Exit\n"
        "}"
    )
    code_run4.font.name = "Courier New"
    code_run4.font.size = Pt(8.5)

    rows_chat_struct = [
        {
            "no": "1", "skenario": "Pengiriman diabaikan karena input kosong atau sedang loading",
            "case": "Simulasi input kosong atau loading aktif",
            "data": "userText = '', loading = true",
            "expected": "Method langsung keluar via early return tanpa memicu state mutation.",
            "actual": "Berhasil sesuai harapan, early return dipicu.", "status": "[√] Valid"
        },
        {
            "no": "2", "skenario": "Pengiriman pesan sukses saat WebSocket stabil",
            "case": "WebSocket dalam status OPEN",
            "data": "userText = 'Konnichiwa', loading = false, readyState = WebSocket.OPEN",
            "expected": "Pesan masuk ke store UI, loading aktif, payload dikirim langsung via WebSocket.",
            "actual": "Berhasil sesuai harapan, pesan terkirim instan.", "status": "[√] Valid"
        },
        {
            "no": "3", "skenario": "Pengiriman pesan ditunda karena WebSocket sedang menghubungkan",
            "case": "WebSocket dalam status CONNECTING",
            "data": "userText = 'Arigatou', loading = false, readyState = WebSocket.CONNECTING",
            "expected": "Pesan masuk ke store UI, listener open dipasang untuk memicu dispatch saat terhubung.",
            "actual": "Berhasil sesuai harapan, listener dipasang.", "status": "[√] Valid"
        }
    ]
    render_table("Tabel 4.17 Skenario Pengujian Struktural sendChat()", rows_chat_struct)

    add_heading_3("2.4.2 Control flow Graph(CFG)")
    add_paragraph(
        "Grafik aliran kontrol (Control Flow Graph) dari method sendChat() dimodelkan untuk mendeteksi "
        "seluruh alur logika percabangan. Kode XML mxGraph (draw.io) berikut merepresentasikan CFG sendChat():"
    )

    xml_p4 = doc.add_paragraph()
    xml_p4.paragraph_format.left_indent = Inches(0.2)
    xml_run4 = xml_p4.add_run(r'''<mxfile host="Electron" modified="2026-07-12T10:00:00.000Z" agent="5.0" version="20.0.0" type="device">
  <diagram id="cfg-send-chat" name="CFG sendChat">
    <mxGraphModel dx="1000" dy="1000" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="827" pageHeight="1169" math="0" shadow="0">
      <root>
        <mxCell id="0" />
        <mxCell id="1" parent="0" />
        <mxCell id="n1" value="Node 1: Entry" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="40" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n2" value="Node 2: if (!userText || loading)" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="160" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n2a" value="Node 2a: return" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="240" y="160" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n3" value="Node 3: chatStore update &amp; loading" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="280" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n4" value="Node 4: if (ws.readyState === OPEN)" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="400" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n5" value="Node 5: dispatch()" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="280" y="520" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n6" value="Node 6: addEventListener open" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="480" y="520" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="n7" value="Node 7: Exit" style="ellipse;whiteSpace=wrap;html=1;aspect=fixed;fillColor=#eeeeee;strokeColor=#333333;" vertex="1" parent="1"><mxGeometry x="380" y="640" width="80" height="80" as="geometry" /></mxCell>
        <mxCell id="e1" parent="1" source="n1" target="n2" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e2" value="Ya" parent="1" source="n2" target="n2a" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e3" value="Tidak" parent="1" source="n2" target="n3" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e4" parent="1" source="n3" target="n4" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e5" value="Ya" parent="1" source="n4" target="n5" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e6" value="Tidak" parent="1" source="n4" target="n6" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e7" parent="1" source="n5" target="n7" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e8" parent="1" source="n6" target="n7" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
        <mxCell id="e9" parent="1" source="n2a" target="n7" edge="1"><mxGeometry relative="1" as="geometry" /></mxCell>
      </root>
    </mxGraphModel>
  </diagram>
</mxfile>''')
    xml_run4.font.name = "Courier New"
    xml_run4.font.size = Pt(7.5)

    add_heading_3("2.4.3 Cyclomatic complexity")
    add_paragraph("Rumus perhitungan Cyclomatic Complexity berdasarkan teori graf adalah sebagai berikut:")

    math_p4 = doc.add_paragraph()
    math_p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math_run4 = math_p4.add_run("V(G) = P + 1")
    math_run4.bold = True
    math_run4.font.name = "Times New Roman"
    math_run4.font.size = Pt(11)

    add_paragraph("Di mana:\n- P = 2 (jumlah simpul keputusan/predicate node pada alur kendali, yaitu Node 2 dan Node 4)")
    add_paragraph("Maka perhitungan nilainya adalah:\n  V(G) = 2 + 1 = 3")
    add_paragraph("Hasil perhitungan menunjukkan nilai Cyclomatic Complexity sebesar 3, menyatakan terdapat 3 jalur logika independen (independent path) utama yang wajib diuji.")

    add_heading_3("2.4.4 Independent path")
    add_paragraph("Berdasarkan perhitungan kompleksitas, rincian dari 3 jalur logika independen (independent path) yang diuji adalah:")
    add_paragraph("Path 1 (1 -> 2 -> 2a -> 7): Aliran eksekusi di mana input query kosong atau sistem sedang memproses pesan lain.")
    add_paragraph("Path 2 (1 -> 2 -> 3 -> 4 -> 5 -> 7): Aliran eksekusi sukses di mana input valid dan WebSocket dalam status OPEN.")
    add_paragraph("Path 3 (1 -> 2 -> 3 -> 4 -> 6 -> 7): Aliran eksekusi di mana input valid namun WebSocket dalam status CONNECTING.")

    add_heading_3("2.4.5 Implementasi dan Hasil Unit testing")
    add_paragraph(
        "Pengujian white-box dengan teknik basis path testing dilakukan menggunakan unit testing otomatis unittest pada "
        "backend dengan memodelkan logic handler Svelte secara presisi. Setiap jalur independen diuji dengan menyimulasikan "
        "masukan state yang berbeda untuk memastikan seluruh percabangan berjalan sesuai spesifikasi melalui perintah "
        "python run_whitebox_sendChat.py."
    )

    rows_chat_unit = [
        {
            "path": "1", "skenario": "Input kosong atau sedang memproses pesan (Guard)",
            "input": "userText = '', loading = true",
            "expected": "Sistem keluar dari method melalui early return seketika",
            "status": "PASS"
        },
        {
            "path": "2", "skenario": "Websocket OPEN pesan dikirim langsung",
            "input": "userText = 'Konnichiwa', loading = false, readyState = WebSocket.OPEN",
            "expected": "Pesan masuk ke store UI, payload terkirim instan via WebSocket",
            "status": "PASS"
        },
        {
            "path": "3", "skenario": "Websocket CONNECTING pengiriman ditunda via listener",
            "input": "userText = 'Arigatou', loading = false, readyState = WebSocket.CONNECTING",
            "expected": "Pesan masuk ke store UI, listener open dipasang sukses",
            "status": "PASS"
        }
    ]
    render_table_5("Tabel 4.18 Hasil Pengujian Unit Testing Method sendChat()", rows_chat_unit)

    add_paragraph(
        "Seluruh 3 jalur independen berhasil dieksekusi dan menghasilkan output sesuai ekspektasi. Hal ini menunjukkan "
        "bahwa validasi masukan pesan chat, pencegahan pengiriman ganda, serta integrasi status WebSocket pada method "
        "sendChat() berjalan secara stabil dan aman dari gangguan konkurensi data. Detail hasil eksekusi ditunjukkan "
        "pada Gambar 4.123 berikut."
    )

    add_paragraph(
        "*(Gambar 4.123 Hasil Unit Testing Method sendChat() ditunjukkan pada tangkapan layar terminal)*",
        italic=True
    )

    add_heading_3("2.4.6 Kesimpulan")
    add_paragraph(
        "Method sendChat() memiliki nilai Cyclomatic Complexity sebesar 3. Seluruh 3 jalur independen berhasil "
        "dieksekusi dengan sukses (PASS) sesuai dengan spesifikasi rancangan logika sistem."
    )

    # 3. Property-Based Testing (PBT)
    add_heading_2("3. Property-Based Testing (PBT)")
    add_paragraph(
        "Property-Based Testing (PBT) menggunakan pustaka Hypothesis untuk memverifikasi invarian "
        "(sifat/hukum sistem yang harus selalu benar) dengan melakukan pengujian fungsionalitas fuzzer otomatis "
        "menggunakan ratusan data acak masukan."
    )
    
    add_paragraph(
        "Secara formal, konsep Property-Based Testing dapat dirumuskan sebagai berikut:"
    )
    
    math_p = doc.add_paragraph()
    math_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math_run = math_p.add_run("∀x ∈ X,    P(x) = True")
    math_run.bold = True
    math_run.font.name = "Times New Roman"
    math_run.font.size = Pt(11)
    
    add_paragraph(
        "Di mana X adalah semesta dari semua data masukan acak (fuzzing inputs) yang digenerasikan secara dinamis "
        "oleh framework pengujian, dan P(x) menyatakan fungsi properti atau invarian yang harus selalu bernilai benar (True) "
        "untuk setiap sampel nilai x."
    )
    t5_rows = [
        {
            "no": "1", "skenario": "P-01: Crash Safety Hiragana",
            "case": "Invarian crash safety NLP konversi Romaji Hiragana",
            "data": "Fuzzing acak Hiragana (~300 sampel)",
            "expected": "Hasil selalu bertipe string & tidak crash",
            "actual": "Hasil valid tipe str (Zero exceptions)", "status": "[√] Valid"
        },
        {
            "no": "2", "skenario": "P-03: Empty Input Invariant",
            "case": "Pencegahan error NLP pada input string kosong/blank",
            "data": "Fuzzing acak spasi & newline (~100 sampel)",
            "expected": "Return string kosong tanpa exception",
            "actual": "Invarian kosong terjaga (No exception)", "status": "[√] Valid"
        },
        {
            "no": "3", "skenario": "P-05: SM-2 EF Floor Invariant",
            "case": "Batas bawah invariant EF harus >= 1.30",
            "data": "Fuzzing initial EF & rating Q (~500 sampel)",
            "expected": "Output EF selalu >= 1.30",
            "actual": "EF >= 1.30 terjaga 100%", "status": "[√] Valid"
        },
        {
            "no": "4", "skenario": "P-06: SM-2 Interval Invariant",
            "case": "Invarian nilai interval review harus berupa bilangan positif",
            "data": "Fuzzing kombinasi Q & EF (~400 sampel)",
            "expected": "interval_days selalu >= 1 hari",
            "actual": "interval_days >= 1 terjaga 100%", "status": "[√] Valid"
        },
        {
            "no": "5", "skenario": "P-07: BKT Boundary Invariant",
            "case": "Invarian nilai probabilitas BKT dalam batas valid",
            "data": "Fuzzing prior probability acak (~500 sampel)",
            "expected": "Output belief selalu dalam range [0.001, 0.999]",
            "actual": "Output belief bounded [0.001, 0.999]", "status": "[√] Valid"
        }
    ]
    render_table("Tabel 4.14 Hasil Pengujian Property-Based Testing (PBT)", t5_rows)

    add_heading_3("3.1 Kesimpulan Pengujian Property-Based Testing (PBT)")
    add_paragraph(
        "Berdasarkan hasil pengujian berbasis sifat (Property-Based Testing) yang dilakukan dengan pustaka "
        "Hypothesis terhadap 5 skenario invarian sistem (mencakup keamanan konversi teks Romaji-Hiragana, "
        "penanganan masukan kosong, pembatasan faktor kemudahan minimum SM-2 EF >= 1.30, interval hari minimum "
        ">= 1, dan pembatasan probabilistik BKT dalam batas [0.001, 0.999]), seluruh pengujian berhasil diselesaikan "
        "dengan status Valid (PASS). Fuzzer otomatis berhasil memverifikasi bahwa invarian sistem tetap terjaga "
        "secara konsisten tanpa menimbulkan kegagalan program (crash), meskipun diberikan ratusan kombinasi data masukan acak "
        "ekstrem. Hal ini membuktikan bahwa modul logika NLP dan algoritma adaptif pada sistem TVJP A.L.I.S.A "
        "memiliki ketahanan (robustness) dan kestabilan data yang sangat tinggi."
    )

    # 4. Metamorphic Testing (MT)
    add_heading_2("4. Metamorphic Testing (MT)")
    add_paragraph(
        "Metamorphic Testing (MT) diimplementasikan untuk memecahkan Oracle Problem pada komponen sistem yang berbasis "
        "non-deterministik/AI. Metode ini mengevaluasi kebenaran keluaran dengan memverifikasi hubungan relasi "
        "metamorfik (Metamorphic Relations atau MR) di antara beberapa eksekusi masukan yang saling terkait."
    )

    add_paragraph(
        "Secara formal, konsep Metamorphic Testing dapat dirumuskan sebagai berikut:"
    )

    math_p = doc.add_paragraph()
    math_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math_run = math_p.add_run("R(f(x),    f(x_new)) = True")
    math_run.bold = True
    math_run.font.name = "Times New Roman"
    math_run.font.size = Pt(11)

    add_paragraph(
        "Di mana f adalah fungsi atau program yang diuji, x adalah data masukan awal (source input), x_new "
        "adalah data masukan lanjutan (follow-up input) yang diperoleh dari hasil transformasi x, dan R menyatakan "
        "hubungan relasi metamorfik (Metamorphic Relation) yang harus selalu bernilai benar (True) untuk "
        "menghubungkan antara keluaran f(x) dan f(x_new)."
    )
    t6_rows = [
        {
            "no": "1", "skenario": "MR-1: BKT Monotonic Increase",
            "case": "Relasi: Prioritas belief harus meningkat jika jawaban benar",
            "data": "f(p, True) vs p",
            "expected": "f(p, True) > p",
            "actual": "Terbukti untuk 10 cases (e.g. 0.7120 > 0.4000)", "status": "[√] Valid"
        },
        {
            "no": "2", "skenario": "MR-1b: BKT Comparative Decrease",
            "case": "Relasi: Output belief jawaban salah harus lebih kecil dari jawaban benar",
            "data": "f(p, False) vs f(p, True)",
            "expected": "f(p, False) < f(p, True)",
            "actual": "Terbukti 11 titik (e.g. 0.2031 < 0.8229)", "status": "[√] Valid"
        },
        {
            "no": "3", "skenario": "MR-2: SM-2 Interval Monotonicity",
            "case": "Relasi: Kualitas jawaban lebih tinggi harus menghasilkan interval lebih panjang",
            "data": "q3 vs q4 pada rep=2",
            "expected": "Interval(q4) >= Interval(q3)",
            "actual": "Terbukti 288 kasus grid (15 >= 12)", "status": "[√] Valid"
        },
        {
            "no": "4", "skenario": "MR-3: SM-2 Reset Invariant",
            "case": "Relasi: Quality kegagalan (Q < 3) harus mereset repetisi ke 0",
            "data": "quality = 0, 1, 2 pada rep=10",
            "expected": "rep == 0 dan interval == 1",
            "actual": "Reset invariant terjaga 45 combo (rep=0, int=1)", "status": "[√] Valid"
        },
        {
            "no": "5", "skenario": "MR-6: SM-2 EF Monotonicity",
            "case": "Relasi: Kualitas jawaban lebih tinggi menghasilkan kenaikan EF lebih besar",
            "data": "q5 vs q4",
            "expected": "EF(q5) >= EF(q4)",
            "actual": "Terbukti 180 pasangan (2.70 >= 2.60)", "status": "[√] Valid"
        }
    ]
    render_table("Tabel 4.15 Hasil Pengujian Metamorphic Testing (MT)", t6_rows)

    add_heading_3("4.1 Kesimpulan Pengujian Metamorphic Testing (MT)")
    add_paragraph(
        "Berdasarkan pengujian metamorfik (Metamorphic Testing) terhadap 5 relasi metamorfik (MR) yang dirancang "
        "untuk mendeteksi deviasi perilaku model kognitif (BKT) dan penjadwalan memori (SM-2), seluruh skenario pengujian "
        "dinyatakan Valid (PASS). Pengujian membuktikan bahwa relasi monotonik BKT (nilai belief meningkat saat jawaban "
        "benar dan menurun saat jawaban salah) serta relasi monotonik SM-2 (kualitas jawaban lebih tinggi menghasilkan "
        "interval review yang lebih panjang/tetap dan nilai EF yang lebih besar) terpenuhi sepenuhnya di ratusan pasangan "
        "data uji. Dengan demikian, pengujian metamorfik ini berhasil memecahkan Oracle Problem pada komponen "
        "kecerdasan buatan (AI) sistem TVJP A.L.I.S.A., memberikan jaminan bahwa mesin pembelajaran adaptif berperilaku "
        "logis dan konsisten secara matematis sesuai dengan teori kognitif."
    )

    # Save
    output_fn = "D_Pengujian_TVJP.docx"
    doc.save(output_fn)
    print(f"Success: Generated {output_fn}")

if __name__ == "__main__":
    create_document()
